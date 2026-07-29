from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from create_business_handbook import setup, table, para, bullet, head, box

OUT=Path(r"D:\FPT\SUMMER_2026\SWP391\API_PAGE_SCOPE_AUTH_PROFILE_CONTRACT_WORKSPACE_COMMON.docx")

def group(d, name, functions, services, routes):
    head(d,name,1)
    para(d,f'Chức năng được giao: {functions}.',lead='Chức năng được giao: ')
    para(d,f'Service source: {services}.',lead='Service source: ')
    para(d,f'Page/route liên quan: {routes}.',lead='Page/route liên quan: ')

def build():
 d=Document(); setup(d)
 q=d.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; q.paragraph_format.space_before=Pt(48); r=q.add_run('AITASKER'); r.bold=True; r.font.name='Arial'; r.font.size=Pt(14); r.font.color.rgb=RGBColor.from_string('2D6BA3')
 q=d.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; q.paragraph_format.space_before=Pt(18); r=q.add_run('API – PAGE THEO\nPHẠM VI PHỤ TRÁCH'); r.bold=True; r.font.name='Arial'; r.font.size=Pt(25); r.font.color.rgb=RGBColor.from_string('173F67')
 q=d.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; q.paragraph_format.space_before=Pt(16); r=q.add_run('Chỉ gồm: Auth • Profile • Contract • Workspace • Common'); r.italic=True; r.font.name='Arial'; r.font.size=Pt(12)
 q=d.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; q.paragraph_format.space_before=Pt(38); r=q.add_run('Không bao gồm Marketplace, Finance, Risk và Admin.'); r.font.name='Arial'; r.font.size=Pt(10.5)
 d.add_page_break()

 head(d,'1. Phạm vi tài liệu',1)
 table(d,['Nhóm','Chức năng giữ lại đúng theo phân công'],[
 ['Auth','Login, Register, OTP, Reset password.'],
 ['Profile','Business KYB, Expert KYC, Portfolio.'],
 ['Contract','Contract list/detail, Sign, NDA, Deposit.'],
 ['Workspace','Milestone, progress report, deliverable, approve/reject.'],
 ['Common','AppShell, Dashboard, Notification, Chatbot.'],
 ])
 box(d,'Quy tắc dùng chung:', 'Các API đi qua apiClient: tự gắn Bearer token, bỏ wrapper `{ success, message, data }` để page nhận `data`, map lỗi thành message. Nếu 401 thì xóa session và điều hướng về /login. Loading được xử lý bằng state/nút disabled ở từng page.')

 group(d,'2. Auth','Login, Register, OTP, Reset password','src/services/authServices.ts','/login • /register • /forgot-password • /reset-password')
 table(d,['Page / thao tác','API – request','Response và dữ liệu hiển thị','Loading / lỗi'],[
 ['LoginPage\n/login – login email','GET /api/auth/check-email?email=…\nPOST /api/auth/login\nbody { email, password }','SessionUser { accessToken, refreshToken, role, accountStatus, email, fullName }. Session lưu token; AppShell hiển thị tên/role.','Nút loading; email/password sai hiện message. 401 xóa session và về login.'],
 ['LoginPage\n/login – Google','POST /api/auth/google/login\nbody { credential, role? }','SessionUser. Nếu là user mới, UI chuyển sang bước chọn role/thông tin để register Google.','Credential/cấu hình Google lỗi: hiện Notice.'],
 ['RegisterPage\n/register – kiểm email/gửi OTP','GET /api/auth/check-email?email=…\nPOST /api/auth/email/send-otp\nbody { email }','check-email trả boolean; send-otp trả { expiresIn } để UI hiển thị bước OTP/thời hạn.','Email trùng chặn đăng ký; đang gửi thì disable.'],
 ['RegisterPage\n/register – xác minh/tạo account','POST /api/auth/email/verify-otp\nbody { email, otp }\nPOST /api/auth/register\nbody { email,password,fullName,phone,role }','SessionUser. Account mới có accountStatus Pending; sau đó user phải làm KYB/KYC.','OTP sai/hết hạn hiện message và giữ form.'],
 ['ForgotPasswordPage\n/forgot-password','GET check-email, sau đó POST /api/auth/forgot-password\nbody { email }','UI báo đã gửi hướng dẫn reset password email.','Email không tồn tại/lỗi backend hiện Notice.'],
 ['ResetPasswordPage\n/reset-password','POST /api/auth/reset-password\nbody { token, newPassword }','Thành công: thông báo và chuyển /login.','Token hết hạn/invalid hoặc password không hợp lệ: hiện error.'],
 ])
 para(d,'Auth API phụ dùng chung: GET /api/auth/me để lấy role/status hiện tại; POST /api/auth/refresh body { refreshToken } để cấp access token mới.')

 group(d,'3. Profile','Business KYB, Expert KYC, Portfolio','src/services/profileService.ts','/app/business/profile • /app/business/kyb • /app/expert/profile • /app/expert/kyc • /app/expert/portfolio')
 table(d,['Page / thao tác','API – request','Response và dữ liệu hiển thị','Loading / lỗi'],[
 ['BusinessProfilePage\n/app/business/profile – load','GET /api/v1/profiles/business/me','BusinessProfile: taxCode, companyName, address, verifiedRepresentative, businessLicenseUrl, kybStatus, rejectionReason. Đổ vào form, badge và preview giấy phép.','Chưa có profile: EmptyState/CTA.'],
 ['BusinessProfilePage/KYB – upload','POST /api/v1/profiles/business/license-file\nFormData { file }','string URL/path. UI gán URL vào businessLicenseUrl để preview/lưu tiếp.','Upload lỗi: Notice, không xóa form.'],
 ['BusinessVerificationProfilePage\n/app/business/kyb – tax check','GET /api/auth/tax-check/:taxCode','TaxCheckResponse { companyName, address, representative, status }. Dùng preview thông tin doanh nghiệp.','taxPreviewLoading; MST lỗi hiện message.'],
 ['Business Profile/KYB – lưu','POST /api/v1/profiles/business\nbody gồm taxCode, businessLicenseUrl và fields profile','BusinessProfile mới với kybStatus/rejectionReason.','Lưu xong cập nhật form/badge; thay đổi hồ sơ có thể quay về Pending để review lại.'],
 ['ExpertProfilePage\n/app/expert/profile – load','GET /api/v1/profiles/expert/me; GET /api/v1/profiles/portfolio/me; GET catalog domain/skill/technology','ExpertProfile + Portfolio; UI đổ nationalId, portfolio URL, years experience, KYC status và các chip.','Các API phụ fallback null/[] để page vẫn mở.'],
 ['Expert Profile/KYC – upload/lưu','POST /api/v1/profiles/expert/portfolio-file\nFormData { file }\nPOST /api/v1/profiles/expert\nbody { nationalId, portfolioUrl, yearsOfExperience }','URL portfolio và ExpertProfile mới/kycStatus.','Thiếu nationalId/portfolioUrl/experience hoặc trùng ID: backend báo lỗi.'],
 ['ExpertPortfolioPage\n/app/expert/portfolio – load/lưu','GET /api/v1/profiles/portfolio/me\nPOST /api/v1/profiles/portfolio\nbody { domainIds,skillIds,technologyIds,yearsExperience,certificates,selfDescription }','Portfolio. UI hiện skills/domain/technology, kinh nghiệm, mô tả và certificate link.','Empty text nếu chưa có dữ liệu; button loading khi save.'],
 ['ExpertPortfolioPage – upload certificate','POST /api/v1/profiles/portfolio/certificate-file\nFormData { file }','string URL/path. Lưu vào certificates để preview/link.','Lỗi upload hiện Notice.'],
 ])
 box(d,'Quy tắc nghiệp vụ Profile:', 'Business KYB yêu cầu tax code đúng và không trùng. Expert KYC yêu cầu nationalId, portfolio URL, years of experience. Khi hồ sơ được gửi/cập nhật, backend có thể đặt account/profile về Pending để chờ xác minh lại.')

 group(d,'4. Contract','Contract list/detail, Sign, NDA, Deposit','src/services/contractService.ts','/app/contracts • /app/contracts/:contractId')
 table(d,['Page / thao tác','API – request','Response và dữ liệu hiển thị','Loading / lỗi'],[
 ['ContractsPage\n/app/contracts – danh sách','GET /api/v1/contracts','Contract[]: contractId, title, business/expert, totalBudget, timelineDays, status, progress, contractMilestones. UI tạo card/list, filter status, link Detail/Workspace.','Code hiện catch lỗi thành []; UI hiện EmptyState “chưa có hợp đồng”.'],
 ['ContractDetailPage\n/app/contracts/:id – load','GET /api/v1/contracts/:contractId\nGET /api/v1/contracts/:contractId/milestones\nGET profile Business/Expert để hiện party','Contract + milestones + profile đối tác. UI hiện điều khoản, ngân sách, tiến độ, chữ ký/NDA và trạng thái.','Không tìm thấy contract: EmptyState. Refresh có loading riêng.'],
 ['Contract Detail – ký','POST /api/v1/contracts/:id/sign','Contract có businessAcceptedAt/expertAcceptedAt/status. UI refresh badge chữ ký.','Nút disabled/loading; sai role/status hiện message backend.'],
 ['Contract Detail – ký NDA','POST /api/v1/contracts/:id/nda-sign','Contract có businessNdaSignedAt/expertNdaSignedAt. Đủ hai chữ ký contract + NDA → status PENDING.','`ndaSubmitting`; lỗi không tự đổi status UI.'],
 ['Contract Detail – Business deposit','POST /api/v1/contracts/:id/deposit/pay','PaymentActionResponse<ContractDeposit>: completed, needTopup, currentBalance, requiredAmount, missingAmount, data. Business giữ 20% total budget.','`depositLoading`; needTopup thì mở/hướng dẫn top-up; success reload contract/wallet.'],
 ['Contract Detail – Expert deposit','POST /api/v1/contracts/:id/expert-deposit/pay','PaymentActionResponse<ContractDeposit>. Expert giữ 10% total budget. Khi hai bên đã giữ tiền, contract PENDING → ACTIVE.','Tương tự Business deposit; không đủ available balance thì không thành công.'],
 ])
 para(d,'Các API khác trong contractService như change request, termination, refund, review không thuộc danh sách chức năng được giao trong ảnh nên không đưa vào phạm vi tài liệu này.')

 group(d,'5. Workspace','Milestone, report, deliverable, approve/reject','src/services/contractService.ts','/app/contracts/:contractId/workspace')
 table(d,['Thao tác trên WorkspacePage','API – request','Response và dữ liệu hiển thị','Ai dùng / trạng thái'],[
 ['Tải workspace','GET /api/v1/contracts/:id\nGET /api/v1/contracts/:id/milestones','Contract + ContractMilestone[]. UI hiện timeline, finalBudget, deadline, status, reject/resubmit count.','Business và Expert; không có contract: EmptyState.'],
 ['Xem tiêu chí/sản phẩm/report của mốc','GET /api/v1/milestones/:id/criteria\nGET /api/v1/milestones/:id/deliverables\nGET /api/v1/contracts/:c/milestones/:m/progress-reports','Criteria[], Deliverable[], MilestoneProgressReport[]. UI hiện checklist, các lần nộp, feedback và lịch sử report.','Dùng khi chọn milestone.'],
 ['Business ký quỹ milestone','POST /api/v1/contracts/:c/milestones/:m/deposit','Milestone status mới. Backend giữ finalBudget trong escrow; UI reload timeline/milestone.','Business; chỉ contract ACTIVE, đúng thứ tự milestone, đủ số dư.'],
 ['Expert bắt đầu','POST /api/v1/milestones/:m/start','Milestone IN_PROGRESS.','Expert; page có nút Start tương thích nếu deposit chưa tự đưa mốc vào IN_PROGRESS.'],
 ['Business yêu cầu report','POST /api/v1/contracts/:c/milestones/:m/progress-report-request','ProgressReportRequestRecord: requestNumber, status, dueAt. UI hiện badge/report deadline.','Business; actionLoading riêng.'],
 ['Expert gửi report','POST /api/v1/contracts/:c/milestones/:m/progress-reports\nbody { content,percentComplete?,attachmentUrl?,sourceCodeUrl?,sourceCodeFileUrl?,demoLink?,submissionNotes? }','MilestoneProgressReport: content, checkpointType, isLate, acknowledgementState, links. UI thêm vào history.','Expert; report trước ACK_PENDING thì backend chặn report tiếp.'],
 ['Business acknowledge/feedback report','POST .../progress-reports/:reportId/acknowledge\nPOST .../feedback body { feedback,category?,severity?,dodItems?,requiresAdjustment? }','Report được ACKNOWLEDGED hoặc có businessFeedback/requiresAdjustment.','Business; UI cập nhật feedback và trạng thái report.'],
 ['Expert upload source/nộp deliverable','POST .../source-code-file FormData { file }\nPOST .../deliverables body { sourceCodeUrl?,sourceCodeFileUrl?,demoLink?,submissionNotes? }','URL source và Deliverable { submissionRound,status,links,notes }. UI hiện chờ review.','Expert; backend kiểm tra deadline, quá hạn không nộp được.'],
 ['Business approve','POST /api/v1/contracts/:c/milestones/:m/approve','Milestone COMPLETED, escrowReleasedAt. Backend chuyển 100% escrow milestone sang Expert.','Business; page reload contract/milestones.'],
 ['Business reject','POST .../reject body { reason, failedCriteria?: [{ criteriaId,reason }] }','Milestone quay về IN_PROGRESS, có rejectCount/resubmitCount/lastRejectionFeedback.','Business; UI hiện reason và tiêu chí fail để Expert sửa.'],
 ])

 group(d,'6. Common','AppShell, Dashboard, Notification, Chatbot','src/layouts/AppLayout/AppLayout.tsx • src/services/notificationService.ts • src/services/chatbotService.ts • src/services/apiClient.ts','AppShell • /app (DashboardPage) • /app/notifications')
 table(d,['Page / component','API – request','Response và dữ liệu hiển thị','Chức năng'],[
 ['AppShell','Dùng SessionUser từ session; token được apiClient tự gắn mọi request.','SessionUser: fullName, role, accountStatus.','Hiện user menu/sidebar theo role; profile gate dựa vào accountStatus.'],
 ['DashboardPage\n/app','Tùy dashboard component; trong phạm vi Common chủ yếu dùng session/layout.','Thông tin user/session và các link điều hướng.','Trang tổng quan sau khi login.'],
 ['NotificationsPage\n/app/notifications – load','GET /api/v1/notifications\nGET /api/v1/notifications/unread-count','NotificationItem[] { type,title,message,targetUrl,isRead,createdAt }; unreadCount cho badge.','Loading hiển thị spinner; [] hiện EmptyState.'],
 ['NotificationsPage – đọc','PATCH /api/v1/notifications/:id/read\nPATCH /api/v1/notifications/read-all','NotificationItem hoặc NotificationItem[] đã cập nhật isRead.','Bấm một notification hoặc “đọc tất cả”; lỗi thì refresh danh sách.'],
 ['AppShell Chatbox','POST /api/chatbot/ask\nbody { question }','ChatbotResponse { answer, sources }. UI thêm question/answer vào lịch sử chat.','Khóa send khi loading; error hiển thị fallback và không mất câu hỏi.'],
 ['Toàn bộ app qua apiClient','Interceptor request/response, không phải endpoint riêng.','`call<T>()` trả `data`; map lỗi.','Gắn Bearer token; 401 clear session + redirect /login; 403 báo không đủ quyền.'],
 ])
 head(d,'7. Câu kết luận để nói khi bảo vệ',1)
 for x in ['“Em phụ trách 5 nhóm: Auth để tạo và bảo vệ session; Profile để xác minh Business/Expert; Contract để ký thỏa thuận và ký quỹ; Workspace để thực thi milestone; Common để điều hướng, thông báo và chatbot.”', '“Mỗi API được khai báo trong service để tái sử dụng, nhưng page mới là nơi điều khiển hành vi UI: loading, hiển thị data, empty state và error.”', '“Các status và số tiền escrow không do frontend tự chỉnh. Sau action, frontend dùng response hoặc reload từ backend để phản ánh đúng dữ liệu.”']:
   bullet(d,x)
 d.core_properties.title='AITASKER - API và Page theo phạm vi Auth Profile Contract Workspace Common'
 d.core_properties.author='AITASKER Team'
 d.save(OUT); print(OUT)

if __name__=='__main__': build()
