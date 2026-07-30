from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from create_business_handbook import setup, table, para, bullet, head, box

OUT=Path(r"D:\FPT\SUMMER_2026\SWP391\API_CATALOG_DAY_DU_AUTH_PROFILE_CONTRACT_WORKSPACE_COMMON.docx")

def heading(d, title, source, pages):
    head(d,title,1); para(d,f'Source frontend: {source}.',lead='Source frontend: '); para(d,f'Page/component dùng: {pages}.',lead='Page/component dùng: ')
def api(d, rows): table(d,['# / chức năng','Endpoint','Input','Response dùng để hiển thị','Page / UI xử lý'], rows)

def build():
 d=Document(); setup(d)
 q=d.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; q.paragraph_format.space_before=Pt(48); r=q.add_run('AITASKER'); r.bold=True; r.font.name='Arial'; r.font.size=Pt(14); r.font.color.rgb=RGBColor.from_string('2D6BA3')
 q=d.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; q.paragraph_format.space_before=Pt(18); r=q.add_run('CATALOG API ĐẦY ĐỦ\nTHEO PHẠM VI ĐƯỢC GIAO'); r.bold=True; r.font.name='Arial'; r.font.size=Pt(24); r.font.color.rgb=RGBColor.from_string('173F67')
 q=d.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; q.paragraph_format.space_before=Pt(16); r=q.add_run('Auth • Profile • Contract • Workspace • Common'); r.italic=True; r.font.name='Arial'; r.font.size=Pt(12)
 q=d.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; q.paragraph_format.space_before=Pt(38); r=q.add_run('Chỉ các endpoint trực tiếp phục vụ những chức năng trong bảng phân công.'); r.font.name='Arial'; r.font.size=Pt(10.5)
 d.add_page_break()

 head(d,'1. Phạm vi và cách đọc',1)
 table(d,['Nhóm','Chức năng trong phạm vi','Không đưa vào'],[
 ['Auth','Login, Register, OTP, Reset password và session bắt buộc cho các chức năng đó.','Marketplace, membership, wallet.'],
 ['Profile','Business KYB, Expert KYC, Portfolio, upload/đọc dữ liệu cần thiết.','Staff verification/review profile.'],
 ['Contract','Danh sách/chi tiết contract, Sign, NDA, Deposit.','Create contract từ proposal, change request, termination, review/refund Admin.'],
 ['Workspace','Milestone, criteria, progress report, source/deliverable, approve/reject.','Dispute, termination, SLA Admin ngoài thao tác thường.'],
 ['Common','AppShell/session, Dashboard shell, notification, chatbot.','PayOS/wallet/finance.'],
 ])
 box(d,'Wrapper chung của mọi API:', 'Backend thường trả `{ success, message, data }`. `src/services/apiClient.ts` unwrap thành `data`; tự gắn `Authorization: Bearer <accessToken>`. `401` xóa session và về `/login`; `403` báo không đủ quyền; timeout báo quá thời gian chờ.')

 heading(d,'2. AUTH – toàn bộ API liên quan Login, Register, OTP, Reset password','src/services/authServices.ts','LoginPage, RegisterPage, ForgotPasswordPage, ResetPasswordPage, session/AppShell')
 api(d,[
 ['1. Kiểm tra email','GET /api/auth/check-email?email={email}','Query: email.','boolean. `true` = email đã tồn tại.','LoginPage kiểm tra account; RegisterPage chặn email trùng; ForgotPasswordPage kiểm tra trước khi gửi email.'],
 ['2. Login thường','POST /api/auth/login','JSON `{ email, password }`.','SessionUser: accessToken, refreshToken, role, accountStatus, email, fullName.','LoginPage lưu session; AppShell dùng tên/role/status. Loading khóa nút; sai thông tin hiện message.'],
 ['3. Google login','POST /api/auth/google/login','JSON `{ credential, role? }`.','SessionUser.','LoginPage/RegisterPage. Nếu Google user mới thì đi sang luồng bổ sung role/thông tin.'],
 ['4. Google register','POST /api/auth/google/register','JSON `{ credential, fullName?, phone, role }`.','SessionUser. Account mới Pending.','LoginPage/RegisterPage lưu session và hướng người dùng làm profile.'],
 ['5. Gửi OTP','POST /api/auth/email/send-otp','JSON `{ email }`.','`{ expiresIn }`.','RegisterPage mở bước OTP, hiển thị/đếm thời gian; disable lúc đang gửi.'],
 ['6. Xác minh OTP','POST /api/auth/email/verify-otp','JSON `{ email, otp }`.','Thành công không cần data UI phức tạp.','RegisterPage chỉ gọi register sau khi OTP đúng; sai/hết hạn hiện error.'],
 ['7. Register thường','POST /api/auth/register','JSON `{ email, password, fullName, phone, role: BUSINESS|EXPERT }`.','SessionUser. Backend tạo account status Pending.','RegisterPage saveSession. Pending nghĩa là phải hoàn thành KYB/KYC.'],
 ['8. Quên password','POST /api/auth/forgot-password','JSON `{ email }`.','Thông báo nghiệp vụ.','ForgotPasswordPage báo đã gửi hướng dẫn; lỗi email hiện Notice.'],
 ['9. Reset password','POST /api/auth/reset-password','JSON `{ token, newPassword }`; token từ URL.','Thông báo thành công.','ResetPasswordPage chuyển về /login; token invalid/hết hạn hiện error.'],
 ['10. Lấy session','GET /api/auth/me','Bearer token.','Partial SessionUser: role, accountStatus, email, fullName.','AppShell/session flow xác định người đang login và quyền hiển thị.'],
 ['11. Refresh token','POST /api/auth/refresh','JSON `{ refreshToken }`.','SessionUser/token mới.','Session flow khi cần token mới; token cũ/không hợp lệ sẽ thành 401.'],
 ])

 heading(d,'3. PROFILE – Business KYB','src/services/profileService.ts','BusinessProfilePage (/app/business/profile), BusinessVerificationProfilePage (/app/business/kyb)')
 api(d,[
 ['1. Lấy hồ sơ Business hiện tại','GET /api/v1/profiles/business/me','Bearer token Business.','BusinessProfile: businessId, taxCode, companyName, address, verifiedRepresentative, businessLicenseUrl, kybStatus, rejectionReason.','Đổ form; StatusBadge; preview giấy phép. Nếu chưa có profile: EmptyState/CTA.'],
 ['2. Tra mã số thuế','GET /api/auth/tax-check/{taxCode}','Path: taxCode.','TaxCheckResponse: taxCode, companyName, address, representative, status.','BusinessVerificationProfilePage preview tự động thông tin doanh nghiệp; có taxPreviewLoading.'],
 ['3. Upload giấy phép','POST /api/v1/profiles/business/license-file','FormData `{ file }`.','string URL/path.','Gán vào `businessLicenseUrl`, preview file, sau đó dùng URL lúc save profile.'],
 ['4. Tạo/cập nhật KYB','POST /api/v1/profiles/business','JSON Partial<BusinessProfile>; cần taxCode, businessLicenseUrl và field form.','BusinessProfile mới: kybStatus/rejectionReason/approvedBy…','Business Profile/KYB cập nhật form/badge. Backend kiểm tra MST 10/13 số, không trùng và có thể đưa hồ sơ về Pending.'],
 ])

 heading(d,'4. PROFILE – Expert KYC và Portfolio','src/services/profileService.ts','ExpertProfilePage (/app/expert/profile), ExpertVerificationProfilePage (/app/expert/kyc), ExpertPortfolioPage (/app/expert/portfolio)')
 api(d,[
 ['1. Lấy Expert profile','GET /api/v1/profiles/expert/me','Bearer token Expert.','ExpertProfile: expertId, nationalId, portfolioUrl, yearsOfExperience, kycStatus, rejectionReason, title/skills…','ExpertProfilePage đổ form và status; ExpertPortfolioPage dùng để ghép thông tin.'],
 ['2. Upload file portfolio','POST /api/v1/profiles/expert/portfolio-file','FormData `{ file }`.','string URL/path.','ExpertProfilePage gán URL vào portfolioUrl, hiển thị file đã chọn.'],
 ['3. Lưu KYC','POST /api/v1/profiles/expert','JSON Partial<ExpertProfile>; cốt lõi `{ nationalId, portfolioUrl, yearsOfExperience }`.','ExpertProfile mới có kycStatus/rejectionReason.','ExpertProfile/KYC page refresh form & badge. National ID không được trùng; gửi lại có thể Pending.'],
 ['4. Lấy portfolio','GET /api/v1/profiles/portfolio/me','Bearer token Expert.','Portfolio: domainIds, skillIds, technologyIds, yearsExperience, certificates, selfDescription.','ExpertProfilePage và ExpertPortfolioPage hiển thị chip kỹ năng/lĩnh vực/công nghệ.'],
 ['5. Upload chứng chỉ','POST /api/v1/profiles/portfolio/certificate-file','FormData `{ file }`.','string URL/path.','ExpertPortfolioPage lưu URL vào certificates, render preview/link.'],
 ['6. Lưu portfolio','POST /api/v1/profiles/portfolio','JSON Partial<Portfolio>: domainIds, skillIds, technologyIds, yearsExperience, certificates, selfDescription.','Portfolio mới.','ExpertPortfolioPage hiển thị lại skills, description, kinh nghiệm và chứng chỉ.'],
 ['7. Dữ liệu danh mục hỗ trợ UI','GET catalog domains/skills/technologies','Query active tùy service catalog.','Danh sách domain/skill/technology.','ExpertProfilePage/PortfolioPage dùng để tạo lựa chọn chips; API phụ lỗi fallback `[]` để page vẫn mở.'],
 ])
 para(d,'Các endpoint Staff review như list/approve Business/Expert không được tính ở đây vì ảnh phân công Profile chỉ nêu Business KYB, Expert KYC và Portfolio.')

 d.add_page_break()
 heading(d,'5. CONTRACT – Contract list/detail, Sign, NDA, Deposit','src/services/contractService.ts','ContractsPage (/app/contracts), ContractDetailPage (/app/contracts/:contractId)')
 api(d,[
 ['1. Danh sách contract','GET /api/v1/contracts','Bearer token.','Contract[]: contractId, jobId, businessId, expertId, contractTitle, totalBudget, timelineDays, status, progress, contractMilestones.','ContractsPage tạo card/list, lọc theo status, link detail/workspace. Lỗi hiện tại fallback `[]` → EmptyState.'],
 ['2. Chi tiết contract','GET /api/v1/contracts/{contractId}','Path: contractId.','Contract: điều khoản, party IDs, totalBudget, timeline, status, chữ ký/NDA timestamps, progress.','ContractDetailPage hiển thị thông tin tổng quan. Không tìm thấy: EmptyState.'],
 ['3. Milestone của contract','GET /api/v1/contracts/{contractId}/milestones','Path: contractId.','ContractMilestone[]: finalBudget, orderIndex, status, criteriaSnapshot, deliverableExpectation, reject/resubmit count, escrow info.','ContractDetailPage hiện milestone summary; WorkspacePage dùng timeline thực thi.'],
 ['4. Profile đối tác bổ sung detail','GET /api/v1/profiles/business/{businessId}; GET /api/v1/profiles/expert/{expertId}','Path: profile ID.','BusinessProfile/ExpertProfile.','ContractDetailPage hiển thị tên/thông tin Business và Expert thay vì chỉ ID.'],
 ['5. Ký contract','POST /api/v1/contracts/{contractId}/sign','Không body.','Contract mới: businessAcceptedAt/expertAcceptedAt/status.','ContractDetailPage refresh badge chữ ký. Chỉ participant đúng role/status mới thành công.'],
 ['6. Ký NDA','POST /api/v1/contracts/{contractId}/nda-sign','Không body.','Contract mới: businessNdaSignedAt/expertNdaSignedAt/status.','ContractDetailPage. Đủ 2 chữ ký contract + 2 NDA → DRAFT sang PENDING.'],
 ['7. Business deposit','POST /api/v1/contracts/{contractId}/deposit/pay','Không body.','PaymentActionResponse<ContractDeposit>: completed, needTopup, currentBalance, requiredAmount, missingAmount, data.','Business giữ 20% totalBudget. UI `depositLoading`; thiếu tiền thì needTopup, không báo thành công.'],
 ['8. Expert deposit','POST /api/v1/contracts/{contractId}/expert-deposit/pay','Không body.','PaymentActionResponse<ContractDeposit>.','Expert giữ 10% totalBudget. Khi hai khoản HELD, backend chuyển contract PENDING → ACTIVE; UI reload contract.'],
 ])

 heading(d,'6. WORKSPACE – tải và quản lý milestone','src/services/contractService.ts','WorkspacePage (/app/contracts/:contractId/workspace)')
 api(d,[
 ['1. Tải contract/workspace','GET /api/v1/contracts/{contractId}','Path: contractId.','Contract: status, party, budget, progress.','WorkspacePage kiểm tra context và render header/tổng quan.'],
 ['2. Tải milestone','GET /api/v1/contracts/{contractId}/milestones','Path: contractId.','ContractMilestone[]/Milestone view.','Timeline: PENDING, IN_PROGRESS, OVERDUE, UNDER_REVIEW, COMPLETED; budget/deadline/reject count.'],
 ['3. Lấy criteria','GET /api/v1/milestones/{milestoneId}/criteria','Path: milestoneId.','AcceptanceCriteria[]: criteriaId, description, sortOrder, isPassed.','WorkspacePage render checklist nghiệm thu cho mốc đang chọn.'],
 ['4. Lấy deliverable','GET /api/v1/milestones/{milestoneId}/deliverables','Path: milestoneId.','Deliverable[]: sourceCodeUrl/file, demoLink, notes, submissionRound, status, rejectionFeedback.','WorkspacePage hiện các lần nộp và lý do reject.'],
 ['5. Ký quỹ milestone','POST /api/v1/contracts/{contractId}/milestones/{milestoneId}/deposit','Không body.','Milestone status mới.','Business bấm Deposit. Backend giữ `finalBudget` trong escrow; kiểm tra contract ACTIVE, thứ tự mốc, số dư.'],
 ['6. Bắt đầu milestone','POST /api/v1/milestones/{milestoneId}/start','Không body.','Milestone IN_PROGRESS.','Expert bấm Start; có thể đã auto-start sau deposit.'],
 ])

 d.add_page_break()
 heading(d,'7. WORKSPACE – Progress report','src/services/contractService.ts','WorkspacePage (/app/contracts/:contractId/workspace)')
 api(d,[
 ['1. Business yêu cầu report','POST /api/v1/contracts/{contractId}/milestones/{milestoneId}/progress-report-request','Không body.','ProgressReportRequestRecord: requestNumber, status, dueAt, progressReportId.','Hiện badge “yêu cầu báo cáo” và deadline.'],
 ['2. Expert nộp report','POST /api/v1/contracts/{contractId}/milestones/{milestoneId}/progress-reports','JSON `{ content, percentComplete?, attachmentUrl?, sourceCodeUrl?, sourceCodeFileUrl?, demoLink?, submissionNotes? }`.','MilestoneProgressReport: checkpointType, content, percent, links, isLate, acknowledgementState.','Thêm report vào history. Nếu report trước ACK_PENDING, backend chặn gửi report tiếp.'],
 ['3. Lấy report','GET /api/v1/contracts/{contractId}/milestones/{milestoneId}/progress-reports','Path: contractId, milestoneId.','MilestoneProgressReport[].','Hiện lịch sử: nội dung, % tiến độ, late, acknowledgement, feedback.'],
 ['4. Acknowledge report','POST .../progress-reports/{reportId}/acknowledge','Không body.','Report có acknowledgementState=ACKNOWLEDGED, acknowledgedAt.','Business xác nhận đã xem.'],
 ['5. Feedback report','POST .../progress-reports/{reportId}/feedback','JSON `{ feedback, category?, severity?, dodItems?, requiresAdjustment? }`.','Report có businessFeedback, feedback category/severity, requiresAdjustment, feedbackAt.','Business góp ý; Expert xem để điều chỉnh.'],
 ])

 d.add_page_break()
 heading(d,'8. WORKSPACE – upload, deliverable, approve/reject','src/services/contractService.ts','WorkspacePage (/app/contracts/:contractId/workspace)')
 api(d,[
 ['1. Upload source ZIP','POST /api/v1/contracts/{contractId}/milestones/{milestoneId}/source-code-file','FormData `{ file }`.','string URL/path.','Expert upload source; UI dùng URL trong report hoặc deliverable. Backend giới hạn ZIP tối đa 50 MB.'],
 ['2. Nộp deliverable','POST /api/v1/contracts/{contractId}/milestones/{milestoneId}/deliverables','JSON `{ sourceCodeUrl?, sourceCodeFileUrl?, demoLink?, submissionNotes? }`.','Deliverable: deliverableId, submissionRound, status=SUBMITTED, links, notes.','Expert nộp sản phẩm; Workspace hiện “chờ Business review”.'],
 ['3. Approve milestone','POST /api/v1/contracts/{contractId}/milestones/{milestoneId}/approve','Không body.','Milestone COMPLETED, escrowReleasedAt, settlement info.','Business duyệt. Backend release 100% escrow mốc vào available balance Expert; UI reload contract/timeline.'],
 ['4. Reject milestone','POST /api/v1/contracts/{contractId}/milestones/{milestoneId}/reject','JSON `{ reason, failedCriteria?: [{ criteriaId, reason }] }`.','Milestone quay IN_PROGRESS, rejectCount/resubmitCount/lastRejectionFeedback; Deliverable REJECTED.','Business bắt buộc nêu reason; UI hiện lý do/criteria fail để Expert sửa và nộp lại.'],
 ])

 d.add_page_break()
 heading(d,'9. COMMON – AppShell, Dashboard, Notification, Chatbot','src/layouts/AppLayout/AppLayout.tsx; src/services/notificationService.ts; src/services/chatbotService.ts; src/services/apiClient.ts','AppShell, DashboardPage (/app), NotificationsPage (/app/notifications), Chatbox')
 api(d,[
 ['1. Session của AppShell','GET /api/auth/me (khi session flow cần đồng bộ)','Bearer token.','Partial SessionUser: role, accountStatus, fullName, email.','AppShell hiện menu/tên/role và profile gate. Dashboard dùng layout/session sau login.'],
 ['2. Lấy notifications','GET /api/v1/notifications','Bearer token.','NotificationItem[]: notificationId, type, title, message, targetUrl, isRead, createdAt, metadata.','NotificationsPage list; AppShell có thể dùng để render dropdown. Loading spinner, rỗng EmptyState.'],
 ['3. Đếm unread','GET /api/v1/notifications/unread-count','Bearer token.','`{ unreadCount }`.','Badge notification ở AppShell.'],
 ['4. Đọc 1 notification','PATCH /api/v1/notifications/{notificationId}/read','Path: notificationId.','NotificationItem đã isRead=true/readAt.','NotificationsPage cập nhật item; lỗi thì refresh list.'],
 ['5. Đọc tất cả','PATCH /api/v1/notifications/read-all','Không body.','NotificationItem[] đã cập nhật.','NotificationsPage bấm “đọc tất cả”, cập nhật list/badge.'],
 ['6. Chatbot','POST /api/chatbot/ask','JSON `{ question }`.','ChatbotResponse `{ answer, sources }`.','AppShell Chatbox append question/answer; loading khóa Send, lỗi hiện fallback.'],
 ['7. Client chung','Không phải endpoint; Axios interceptor/call<T>.','Tự gắn token, tự unwrap response.','`data` đã lấy khỏi wrapper, error được map message.','Mọi page trong 5 nhóm dùng cơ chế này; 401 redirect login, 403 thông báo không quyền.'],
 ])
 head(d,'10. Câu trả lời chốt khi bảo vệ',1)
 for x in ['“Tổng API trong tài liệu này chỉ là API phục vụ 5 nhóm chức năng được giao. Các API Marketplace, Finance, Risk và Admin đã được loại ra để không lẫn phạm vi.”', '“Service khai báo endpoint; page gọi endpoint theo hành động UI. Response quan trọng được đưa lên form, badge status, timeline milestone, danh sách notification hoặc chat history.”', '“Frontend không tự quyết định quyền, status hay tiền. Nó gửi request, hiển thị loading/error và dùng response/reload từ backend làm nguồn dữ liệu đúng.”']:
   bullet(d,x)
 d.core_properties.title='AITASKER - Catalog API day du theo pham vi duoc giao'
 d.core_properties.author='AITASKER Team'
 d.save(OUT); print(OUT)

if __name__=='__main__': build()
