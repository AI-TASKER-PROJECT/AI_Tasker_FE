from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(r"D:\FPT\SUMMER_2026\SWP391\SO_TAY_NGHIEP_VU_AITASKER_CHI_TIET.docx")
NAVY, BLUE, PALE, NOTE, GREY = "173F67", "2D6BA3", "F4F8FC", "E8F1FA", "5B6775"

def cell(c, fill=None):
    tc=c._tc.get_or_add_tcPr(); mar=OxmlElement('w:tcMar')
    for s,v in [('top',100),('start',110),('bottom',100),('end',110)]:
        x=OxmlElement(f'w:{s}'); x.set(qn('w:w'),str(v)); x.set(qn('w:type'),'dxa'); mar.append(x)
    tc.append(mar)
    if fill:
        x=OxmlElement('w:shd'); x.set(qn('w:fill'),fill); tc.append(x)
    c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER

def table(d, heads, rows):
    t=d.add_table(rows=1,cols=len(heads)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.LEFT
    for i,x in enumerate(heads):
        c=t.rows[0].cells[i]; cell(c,NAVY); r=c.paragraphs[0].add_run(x); r.bold=True; r.font.name='Arial'; r.font.size=Pt(8.7); r.font.color.rgb=RGBColor(255,255,255)
    for n,row in enumerate(rows):
        cs=t.add_row().cells
        for i,x in enumerate(row):
            c=cs[i]; cell(c,PALE if n%2 else None); p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.03
            r=p.add_run(x); r.font.name='Arial'; r.font.size=Pt(8.55)
        t.rows[-1]._tr.get_or_add_trPr().append(OxmlElement('w:cantSplit'))
    d.add_paragraph().paragraph_format.space_after=Pt(2)

def para(d,text,lead=None):
    q=d.add_paragraph(); q.paragraph_format.space_after=Pt(5); q.paragraph_format.line_spacing=1.1
    if lead and text.startswith(lead):
        r=q.add_run(lead); r.bold=True; r.font.name='Arial'; r.font.size=Pt(10.4); text=text[len(lead):]
    r=q.add_run(text); r.font.name='Arial'; r.font.size=Pt(10.4)
def bullet(d,text):
    q=d.add_paragraph(style='List Bullet'); q.paragraph_format.space_after=Pt(3)
    r=q.add_run(text); r.font.name='Arial'; r.font.size=Pt(10.1)
def head(d,text,level=1): d.add_paragraph(text,style=f'Heading {level}')
def box(d,title,text):
    t=d.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.LEFT; c=t.cell(0,0); cell(c,NOTE)
    q=c.paragraphs[0]; q.paragraph_format.space_after=Pt(2); r=q.add_run(title); r.bold=True; r.font.name='Arial'; r.font.size=Pt(10.1); r.font.color.rgb=RGBColor.from_string(NAVY)
    q=c.add_paragraph(); q.paragraph_format.space_after=Pt(1); r=q.add_run(text); r.font.name='Arial'; r.font.size=Pt(10.0)
    t.rows[0]._tr.get_or_add_trPr().append(OxmlElement('w:cantSplit'))
    d.add_paragraph().paragraph_format.space_after=Pt(2)

def setup(d):
    s=d.sections[0]; s.top_margin=Inches(.66); s.bottom_margin=Inches(.62); s.left_margin=Inches(.7); s.right_margin=Inches(.7)
    n=d.styles['Normal']; n.font.name='Arial'; n._element.rPr.rFonts.set(qn('w:ascii'),'Arial'); n._element.rPr.rFonts.set(qn('w:hAnsi'),'Arial'); n.font.size=Pt(10.4)
    for name,size,color in [('Heading 1',16,NAVY),('Heading 2',12.5,BLUE),('Heading 3',11,NAVY)]:
        st=d.styles[name]; st.font.name='Arial'; st._element.rPr.rFonts.set(qn('w:ascii'),'Arial'); st._element.rPr.rFonts.set(qn('w:hAnsi'),'Arial'); st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color); st.paragraph_format.space_before=Pt(12); st.paragraph_format.space_after=Pt(6)
    h=s.header.paragraphs[0]; h.alignment=WD_ALIGN_PARAGRAPH.RIGHT; r=h.add_run('AITASKER | SỔ TAY NGHIỆP VỤ CHI TIẾT'); r.font.name='Arial'; r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(GREY)
    f=s.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=f.add_run('Đã đối chiếu source frontend và backend | Auth • Profile • Contract • Workspace • Common'); r.font.name='Arial'; r.font.size=Pt(7.6); r.font.color.rgb=RGBColor.from_string(GREY)

def build():
 d=Document(); setup(d)
 q=d.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; q.paragraph_format.space_before=Pt(46); r=q.add_run('AITASKER'); r.bold=True; r.font.name='Arial'; r.font.size=Pt(14); r.font.color.rgb=RGBColor.from_string(BLUE)
 q=d.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; q.paragraph_format.space_before=Pt(18); r=q.add_run('SỔ TAY NGHIỆP VỤ\nCHI TIẾT TỪ CODE'); r.bold=True; r.font.name='Arial'; r.font.size=Pt(25); r.font.color.rgb=RGBColor.from_string(NAVY)
 q=d.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; q.paragraph_format.space_before=Pt(14); r=q.add_run('Dành cho người chưa biết code – phạm vi Auth, Profile, Contract, Workspace, Common'); r.italic=True; r.font.name='Arial'; r.font.size=Pt(11.5); r.font.color.rgb=RGBColor.from_string(GREY)
 q=d.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; q.paragraph_format.space_before=Pt(40); r=q.add_run('Mục tiêu: hiểu dự án đang giải quyết vấn đề gì, mỗi vai trò làm gì, tiền đi đâu, trạng thái thay đổi thế nào, page/API nào thực hiện từng bước.'); r.font.name='Arial'; r.font.size=Pt(10.5)
 d.add_page_break()

 head(d,'1. Dự án này thực sự làm gì?',1)
 para(d,'AITASKER là nền tảng trung gian giúp Business thuê Expert thực hiện công việc theo từng mốc (milestone). Hệ thống không chỉ “đăng job và chat”; nó kiểm soát quy trình để giảm rủi ro: xác minh người dùng, chốt thỏa thuận, giữ tiền, theo dõi tiến độ, nghiệm thu và xử lý tranh chấp.')
 table(d,['Ai','Họ cần gì?','Hệ thống bảo vệ điều gì?'],[
 ['Business','Thuê Expert, có sản phẩm đúng yêu cầu và không mất tiền khi chưa có kết quả.','Chỉ giải ngân tiền milestone sau khi Business nghiệm thu hoặc theo SLA.'],
 ['Expert','Nhận việc rõ ràng, có tiền đảm bảo và được thanh toán khi hoàn thành.','Tiền milestone được Business ký quỹ trước khi Expert làm.'],
 ['Staff','Kiểm tra danh tính/hồ sơ và hỗ trợ xử lý case.','Chỉ Staff thuộc domain phù hợp mới được review profile/tranh chấp.'],
 ['Admin','Quản trị quy tắc, tài chính và các case cấp hệ thống.','Quyết định/thi hành settlement, hoàn ký quỹ, quản trị cấu hình.'],
 ])
 head(d,'Luồng nghiệp vụ lớn nhất',2)
 box(d,'Hãy hình dung như thuê thợ làm nhà:', 'Tạo tài khoản → xác minh “có thật/đủ năng lực” → Business chọn Expert → lập hợp đồng → hai bên ký và đặt tiền bảo đảm → Business ký quỹ tiền của từng hạng mục → Expert làm/nộp sản phẩm → Business nghiệm thu → hệ thống trả tiền cho Expert. Nếu có bất đồng: tự thương lượng → Staff/Admin can thiệp → hệ thống chia tiền theo quyết định.')
 para(d,'Chuỗi trạng thái chính: Account Pending → Approved → Contract DRAFT → PENDING → ACTIVE → COMPLETED/TERMINATED/CANCELLED. Mỗi milestone: PENDING → IN_PROGRESS → UNDER_REVIEW → COMPLETED; hoặc quay lại IN_PROGRESS khi bị reject; hoặc DISPUTED khi có tranh chấp.')

 head(d,'2. Kiến trúc dễ hiểu: click ở đâu thì chuyện gì xảy ra?',1)
 table(d,['Lớp','Nơi trong dự án','Vai trò'],[
 ['Giao diện (FE)','Page/Component ở AI_Tasker_FE','Hiện form, nút, bảng; gọi API và hiển thị loading/lỗi/kết quả.'],
 ['Service FE','src/services','Đóng gói endpoint như login(), getContract(), approveMilestone().'],
 ['API backend','Controller ở AI_Tasker_BE','Nhận request theo URL và chuyển vào service.'],
 ['Nghiệp vụ backend','Service ở AI_Tasker_BE','Kiểm tra role, trạng thái, điều kiện tiền; thay đổi database; tạo notification/audit log.'],
 ['Database/ví','Entity/Repository/WalletLedger','Lưu trạng thái; số tiền thật được hạch toán bằng ledger/escrow, không tin dữ liệu frontend.'],
 ])
 box(d,'Cách trả lời khi bị hỏi một nút bất kỳ:', '“Nút ở route … gọi API …; backend kiểm tra role và trạng thái. Nếu hợp lệ, backend cập nhật dữ liệu rồi trả response; frontend cập nhật UI. Nếu lỗi, UI chỉ báo message, không tự đổi trạng thái.”')

 head(d,'3. Auth – ai được vào hệ thống và vào như thế nào?',1)
 para(d,'Auth không chỉ kiểm tra mật khẩu. Nó tạo phiên đăng nhập, biết role, khóa tài khoản khi có dấu hiệu tấn công, và làm cho token cũ mất hiệu lực khi có lần login mới.')
 table(d,['Page / route','Người dùng làm gì','API','Backend thực sự làm gì'],[
 ['Đăng ký\n/register','Chọn BUSINESS hoặc EXPERT; xác minh email OTP; gửi form.','GET /api/auth/check-email\nPOST /api/auth/email/send-otp\nPOST /api/auth/email/verify-otp\nPOST /api/auth/register','Chỉ register nếu email đã OTP verified và chưa tồn tại. Account mới có status Pending, emailVerified=true; được tạo quota ví ban đầu và Admin nhận notification.'],
 ['Đăng nhập\n/login','Nhập email/password hoặc Google.','POST /api/auth/login\nPOST /api/auth/google/login','Kiểm tra password, tài khoản bị khóa tạm/vĩnh viễn. Sai nhiều lần: có thể khóa tạm; lặp lại sau đó có thể khóa bảo mật và yêu cầu reset password.'],
 ['Lấy lại session','Mở app hoặc token access hết hạn.','GET /api/auth/me\nPOST /api/auth/refresh','Trả role, accountStatus, email, fullName; refresh token chỉ hợp lệ nếu đúng token version hiện tại.'],
 ['Quên/đổi password','Nhập email rồi dùng link reset.','POST /api/auth/forgot-password\nPOST /api/auth/reset-password','Xử lý token reset; sau khi đổi mật khẩu, thông báo bảo mật và phiên cũ không còn là nguồn tin cậy.'],
 ])
 head(d,'Session trả về dùng để làm gì?',2)
 table(d,['Trường trả về','Frontend dùng ở đâu','Ý nghĩa nghiệp vụ'],[
 ['accessToken','Axios Authorization: Bearer <token>','Backend nhận diện request này là của ai.'],
 ['refreshToken','Khi cần lấy access token mới','Không bắt người dùng đăng nhập lại ngay khi access token hết hạn.'],
 ['role','Sidebar, RoleProtectedRoute, nút hành động','Chỉ hiển thị luồng phù hợp Business/Expert/Staff/Admin.'],
 ['accountStatus','AppShell/profile gate','Pending/Rejected không được đi sâu vào marketplace hoặc contract execution.'],
 ['fullName/email','User menu và thông tin tài khoản','Hiển thị danh tính đang dùng app.'],
 ])
 bullet(d,'Loading: khóa nút submit để không tạo hai request. Sai password/OTP: hiện message, giữ form. 401: clear session và về /login. 403: báo không đủ quyền.')
 bullet(d,'Điểm bảo mật để nói: login mới tăng activeTokenVersion; token cũ không khớp version bị từ chối. Frontend chỉ hỗ trợ UX; backend mới là nơi quyết định quyền cuối cùng.')

 head(d,'4. Profile / KYC-KYB – vì sao đăng ký xong vẫn chưa làm được việc?',1)
 para(d,'Sau register, tài khoản là Pending. Điều này có nghĩa “đã tạo account, nhưng chưa được tin cậy để giao dịch”. Business cần KYB (xác minh doanh nghiệp), Expert cần KYC (xác minh chuyên gia).')
 table(d,['Vai trò / page','Route và API','Quy tắc backend xác nhận'],[
 ['Business profile','/app/business/profile\nGET /api/v1/profiles/business/me\nPOST /api/v1/profiles/business','Tax code bắt buộc, phải là 10 hoặc 13 số, không được trùng account khác. Backend gọi tax-check để lấy companyName, address, representative; không tin hoàn toàn dữ liệu người dùng tự gõ.'],
 ['Business license','/app/business/kyb\nPOST /api/v1/profiles/business/license-file\nGET /api/auth/tax-check/:taxCode','Upload file để có URL/path preview. Khi gửi/cập nhật KYB theo submission mode, profile và account quay lại Pending để Staff review lại.'],
 ['Expert profile','/app/expert/profile\nGET /api/v1/profiles/expert/me\nPOST /api/v1/profiles/expert','National ID, portfolio URL, years of experience bắt buộc; National ID không được trùng. Cập nhật KYC có thể đưa account về Pending để duyệt lại.'],
 ['Portfolio','/app/expert/portfolio\nGET /api/v1/profiles/portfolio/me\nPOST /api/v1/profiles/portfolio\nPOST /api/v1/profiles/portfolio/certificate-file','Lưu kinh nghiệm/năng lực/chứng chỉ. Backend đồng bộ số năm kinh nghiệm từ Expert Profile sang portfolio.'],
 ['Review','/app/verifications\nPOST /api/v1/profiles/approve/{type}/{id}?status=...&reason=...','Code backend hiện tại yêu cầu role STAFF cho thao tác duyệt. Chỉ profile đang Pending mới review; chỉ Approved hoặc Rejected hợp lệ; Rejected cần lý do.'],
 ])
 box(d,'Trạng thái profile có ý nghĩa gì?', 'Pending: chờ Staff kiểm tra. Approved: tài khoản được phép đi vào các nghiệp vụ Business/Expert yêu cầu “approved account”. Rejected: bị từ chối và có rejection reason để sửa/nộp lại. Lock: không thể nhận token/login bình thường.')
 para(d,'Dữ liệu được hiển thị: form đổ companyName/taxCode/license hoặc nationalId/portfolio/experience; badge hiển thị KYC/KYB/account status; Staff xem file qua API file view URL; danh sách verification lọc hồ sơ chờ xử lý.')

 head(d,'5. Hợp đồng: từ proposal đến dự án đang hoạt động',1)
 para(d,'Contract được tạo từ một proposal đã Accepted. Nó “đóng băng” các mốc công việc và ngân sách đã chốt, để về sau Job thay đổi không làm mơ hồ thỏa thuận cũ.')
 table(d,['Bước','API / người làm','Backend kiểm tra và kết quả'],[
 ['1. Tạo hợp đồng nháp','POST /api/v1/contracts/from-proposals/{proposalId}\nBusiness','Business phải Approved; proposal phải Accepted, chưa từng tạo contract; job phải thuộc chính Business; job phải có milestone. Contract tạo ở DRAFT.'],
 ['2. Chụp snapshot milestone','Tự chạy trong bước tạo','Mỗi job milestone được copy thành contract milestone: tên, mô tả, thứ tự, ngân sách gốc/final, duration, acceptance criteria snapshot, trạng thái PENDING.'],
 ['3. Hai bên ký contract','POST /api/v1/contracts/{id}/sign\nBusiness và Expert','Chỉ participant và đã Approved được ký. Backend lưu businessAcceptedAt/expertAcceptedAt.'],
 ['4. Hai bên ký NDA','POST /api/v1/contracts/{id}/nda-sign\nBusiness và Expert','Backend lưu businessNdaSignedAt/expertNdaSignedAt. Chỉ khi đủ 4 dấu mốc: 2 chữ ký contract + 2 chữ ký NDA thì contract chuyển DRAFT → PENDING.'],
 ['5. Hai bên ký quỹ hợp đồng','POST /api/v1/contracts/{id}/deposit/pay (Business)\nPOST /api/v1/contracts/{id}/expert-deposit/pay (Expert)','Business giữ 20% tổng budget; Expert giữ 10% tổng budget trong escrow. Không đủ available balance: response needTopup=true/chỉ rõ thiếu bao nhiêu. Khi cả hai HELD, contract PENDING → ACTIVE.'],
 ])
 head(d,'Trạng thái Contract – nói sao cho đúng?',2)
 table(d,['Status','Nghĩa','UI nên làm gì'],[
 ['DRAFT','Có hợp đồng nháp, chưa đủ chữ ký/NDA.','Hiện nút ký hoặc ký NDA tùy người dùng chưa làm gì; Business chỉ cancel draft khi chưa ai ký/NDA.'],
 ['PENDING','Đủ chữ ký và NDA, đang chờ hai bên ký quỹ contract.','Hiện trạng thái ai đã/chưa ký quỹ; thiếu tiền thì dẫn top-up.'],
 ['ACTIVE','Đủ hai khoản ký quỹ; được thực hiện milestone.','Mở Workspace và action ký quỹ từng milestone.'],
 ['COMPLETED','Tất cả milestone completed.','Đóng job thành CLOSED, tự hoàn ký quỹ hợp đồng cho các bên.'],
 ['CANCELLED','Expert từ chối hoặc Business hủy nháp hợp lệ.','Job trở lại OPEN/review proposal theo trường hợp.'],
 ['TERMINATION_PENDING / TERMINATED','Đang/đã chấm dứt.','Chặn luồng thực thi thường, theo dõi termination/settlement.'],
 ])
 para(d,'Page liên quan: /app/contracts gọi GET /api/v1/contracts để hiện danh sách (budget, status, progress, party). /app/contracts/:contractId gọi GET /api/v1/contracts/:id và GET /api/v1/contracts/:id/milestones để hiện điều khoản, chữ ký/NDA/ký quỹ và các mốc.')

 head(d,'6. Tiền trong hệ thống: có 2 lớp “giữ tiền”',1)
 para(d,'Đây là phần rất quan trọng để phân biệt: “ký quỹ hợp đồng” và “escrow của milestone” không phải một.')
 table(d,['Loại tiền','Ai trả','Khi nào?','Mục đích / kết thúc'],[
 ['Ký quỹ hợp đồng','Business 20% tổng budget; Expert 10% tổng budget','Sau khi contract đủ chữ ký/NDA và đang PENDING.','Tiền bảo đảm cam kết. Giữ HELD; khi contract hoàn thành thì tự hoàn. Khi chấm dứt, hoàn/khấu trừ theo settlement/quy tắc.'],
 ['Escrow milestone','Business trả đúng finalBudget của milestone','Trước khi Expert làm từng milestone, theo đúng orderIndex.','Bảo đảm tiền làm việc. Backend hold từ available balance. Khi Business approve, debit escrow và credit toàn bộ cho Expert.'],
 ['Top-up ví','Người dùng nạp qua PayOS','Khi available balance không đủ.','PayOS create/sync; frontend chỉ reload số dư sau khi backend xác nhận PAID, không tự cộng tiền.'],
 ])
 box(d,'Ví dụ dễ nói:', 'Contract có tổng 100 triệu: Business ký quỹ hợp đồng 20 triệu, Expert ký quỹ 10 triệu. Milestone 1 trị giá 30 triệu: trước khi Expert làm, Business phải đưa 30 triệu từ available balance vào escrow milestone. Khi Business duyệt, 30 triệu đó mới được chuyển sang available balance của Expert. Đây là lý do Expert có bảo đảm thanh toán nhưng Business vẫn kiểm soát nghiệm thu.')

 head(d,'7. Workspace: vòng đời một milestone từ đầu đến cuối',1)
 para(d,'Route: /app/contracts/:contractId/workspace. Đây là nơi thực thi thực tế. Backend buộc Business làm milestone theo thứ tự orderIndex: chưa hoàn tất mốc trước thì không được ký quỹ mốc sau.')
 table(d,['Trạng thái','Ai / API','Backend làm gì','UI hiển thị gì'],[
 ['PENDING','Business: POST /api/v1/contracts/{contractId}/milestones/{milestoneId}/deposit','Kiểm tra contract ACTIVE, đúng chủ contract, mốc chưa deposit, mốc trước đã completed. Hold finalBudget vào escrow; trong code hiện tại deposit có thể đưa mốc trực tiếp IN_PROGRESS.','Nút “Ký quỹ/Bắt đầu”; ngân sách, deadline, thứ tự mốc.'],
 ['IN_PROGRESS','Expert: POST /api/v1/milestones/{milestoneId}/start (tương thích); POST .../progress-reports','Chỉ Expert Approved; chỉ mốc đang làm hoặc quá hạn. Lưu report, % tiến độ, attachment/source/demo link; report mới cần Business acknowledge trước report tiếp.','Timeline đang làm, form báo cáo, lịch sử report, badge “chờ Business xác nhận”.'],
 ['OVERDUE','Admin: POST /api/v1/contracts/{id}/milestones/check-overdue','Nếu quá deadline dựa trên inProgressStartedAt + duration, backend đánh dấu OVERDUE và thông báo hai bên. Expert vẫn có thể gửi progress report; deliverable bị chặn khi đã quá hạn.','Cảnh báo quá hạn, không cho Expert nộp deliverable muộn.'],
 ['UNDER_REVIEW','Expert: POST .../deliverables; hoặc source-code-file; complete flow','Nộp deliverable/source code. Backend kiểm tra deadline, participant và trạng thái. Sản phẩm chờ Business nghiệm thu.','Link/tệp, submission round, thời điểm nộp, nút Approve/Reject cho Business.'],
 ['COMPLETED','Business: POST .../approve','Chỉ Business Approved, milestone UNDER_REVIEW. Backend release escrow: trừ tiền đang giữ của Business và cộng 100% finalBudget cho Expert; đánh dấu escrow released. Nếu mọi mốc completed thì contract COMPLETED.','Thông báo thành công, milestone hoàn thành, số tiền được giải ngân.'],
 ['Quay lại IN_PROGRESS','Business: POST .../reject, body reason + failedCriteria','Reason là bắt buộc. Deliverable hiện tại thành REJECTED; reject/resubmit count tăng; mốc trở lại IN_PROGRESS để Expert sửa và nộp lại.','Hiện lý do reject, từng tiêu chí không đạt, số lần resubmit.'],
 ['DISPUTED','Business hoặc Expert: POST .../disputes','Chỉ ACTIVE và mốc IN_PROGRESS/OVERDUE/UNDER_REVIEW; không được có dispute hoặc termination đang active. Mốc bị đóng băng ở DISPUTED.','Badge tranh chấp; chặn hành động thường, mở lịch sử/evidence.'],
 ])
 head(d,'Progress report có quy tắc gì?',2)
 bullet(d,'Business có thể request report. Request thứ nhất có hạn 24 giờ; các lần sau 12 giờ. Nếu request cũ quá hạn, backend đánh dấu expired trước khi tạo request mới.')
 bullet(d,'Expert chỉ gửi report khi milestone IN_PROGRESS/OVERDUE. Báo cáo phải có nội dung; backend gắn checkpoint MIDPOINT/PRE_DEADLINE khi phù hợp, isLate nếu trễ.')
 bullet(d,'Business acknowledge hoặc feedback. Nếu report gần nhất còn ACK_PENDING, Expert không được nộp report tiếp — tránh tạo nhiều báo cáo mà chưa ai xem.')

 head(d,'8. Nếu không đồng ý: change request, dispute, termination',1)
 table(d,['Tình huống','Luồng nghiệp vụ thực tế','Điều cần nói khi bảo vệ'],[
 ['Muốn đổi phạm vi/ngân sách/thời gian','POST/GET /api/v1/contracts/{id}/change-requests; accept/reject. Khi accept, backend áp dụng proposed budget/timeline/scope/milestones. Chỉ milestone PENDING và escrow chưa release mới được sửa.','Không cho sửa mốc đã bắt đầu/đã giải ngân, vì điều đó phá vỡ cam kết và lịch sử tiền.'],
 ['Tranh chấp milestone','Tạo dispute → PENDING_SELF_RESOLVE → một bên yêu cầu escalation → ESCALATION_REQUESTED → Staff route/auto-route → STAFF_REVIEWING → Staff decision → Admin execute settlement.','Trước tiên hai bên tự giải quyết. Khi escalated, Staff có thời hạn thu thập evidence 48h và SLA review; settlement mới xử lý tiền escrow.'],
 ['Chấm dứt contract','Business/Expert tạo termination request; bên còn lại phản hồi hoặc dispute; Staff review/approve; Admin execute settlement; hoàn ký quỹ.','Termination chặn việc thực thi thường. Settlement tính Expert payout % và Business refund từ escrow mốc hiện tại; các mốc còn lại bị cancel.'],
 ])
 para(d,'Ví dụ termination: nếu Staff/Admin xác định Expert được nhận 40% tiền milestone đang giữ, backend debit escrow, credit Expert 40%, credit Business 60%, đánh dấu mốc completed hoặc cancelled tùy payout, hủy các mốc sau, rồi hoàn ký quỹ contract theo flow.')

 head(d,'9. Common: notification, chatbot, PayOS và UI state',1)
 table(d,['Chức năng','API / dữ liệu','Nghiệp vụ và UI'],[
 ['Notification','GET /api/v1/notifications\nGET /api/v1/notifications/unread-count\nPATCH .../:id/read\nPATCH .../read-all','Backend tạo notification ở nhiều điểm: profile submitted/reviewed, contract created/signed, deposit held, report submitted, milestone approved/rejected/overdue, dispute/termination. UI badge là unread count; list rỗng hiện empty state.'],
 ['Chatbot','POST /api/chatbot/ask → answer, sources','UI thêm answer/sources vào chat history. Khi loading khóa gửi; khi lỗi hiện fallback và giữ lại câu hỏi. Chatbot không phải nguồn quyết định nghiệp vụ/tiền.'],
 ['Top-up PayOS','POST /api/payments/payos/create\nPOST /api/payments/payos/{orderCode}/sync\nGET /api/v1/wallet/me','Create cho checkout URL/QR/orderCode. Sync để backend kiểm tra trạng thái. UI chỉ refresh wallet sau response xác nhận, không tự cộng tiền tại client.'],
 ['API states chung','ApiResponse { success, message, data }','Loading: spinner/disable. Success: dùng data trả về để refresh state. Empty: CTA đúng ngữ cảnh. Error: map message, giữ form. 401: logout; 403: không đủ quyền.'],
 ])

 head(d,'10. Ma trận quyền đúng theo code backend',1)
 table(d,['Hành động','Business','Expert','Staff','Admin'],[
 ['Đăng ký / login','Có (account mới Pending).','Có (account mới Pending).','Có account có sẵn/quản trị cấp.','Có account quản trị.'],
 ['Khai profile','Tạo KYB, upload license.','Tạo KYC/portfolio/upload chứng chỉ.','Không tạo KYB/KYC cho mình trong luồng này.','Quản trị theo API riêng.'],
 ['Duyệt profile','Không.','Không.','Có: code ProfileService yêu cầu STAFF, domain phù hợp; approve/reject Pending profile.','Không phải endpoint approve profile hiện tại trừ khi có luồng/API riêng.'],
 ['Tạo contract từ proposal','Có, nếu proposal Accepted, job của mình, profile Approved.','Không.','Không.','Không phải luồng thường.'],
 ['Ký contract/NDA','Có nếu là party và Approved.','Có nếu là party và Approved.','Không.','Không.'],
 ['Ký quỹ contract','20% tổng budget.','10% tổng budget.','Không.','Có thể refund qua API Admin.'],
 ['Ký quỹ milestone / approve/reject','Deposit và approve/reject.','Không.','Không.','Chạy SLA/check overdue theo API Admin.'],
 ['Làm/nộp report/deliverable','Xem, feedback/ack report.','Start, report, upload/nộp sản phẩm.','Xem case được giao.','Quản trị.'],
 ['Dispute / settlement','Tạo, escalation, evidence.','Tạo, escalation, evidence.','Route/review/decision nếu được assign.','Quản trị/execute settlement, termination settlement.'],
 ])

 head(d,'11. Kịch bản demo 8 phút và câu hỏi thường bị hỏi',1)
 table(d,['Phút','Demo','Câu nói gợi ý'],[
 ['0–1','Login Business; chỉ role/accountStatus.','“Token và role giúp backend biết request của ai. Account Pending chỉ mới đăng ký, chưa được giao dịch.”'],
 ['1–2','Mở Profile/KYB và verification status.','“Hồ sơ là lớp tin cậy. Tax code được backend kiểm tra; Staff duyệt Pending thành Approved/Rejected.”'],
 ['2–4','Mở Contract Detail.','“Contract chỉ Active sau 2 chữ ký contract, 2 NDA và 2 khoản ký quỹ: Business 20%, Expert 10%.”'],
 ['4–6','Mở Workspace milestone.','“Business ký quỹ final budget của milestone; tiền bị giữ escrow. Expert gửi report/nộp deliverable; Business approve mới chuyển 100% tiền cho Expert.”'],
 ['6–7','Minh họa reject hoặc dispute.','“Reject cần reason và failed criteria, đưa mốc về In Progress. Dispute đóng băng mốc để không giải ngân sai.”'],
 ['7–8','Notifications/top-up.','“Mọi mốc chính đều tạo notification. Top-up chỉ tăng số dư khi backend sync PayOS xác nhận.”'],
 ])
 for q,a in [
 ('Tại sao Business phải ký quỹ cả contract và milestone?', 'Ký quỹ contract là tiền bảo đảm cam kết của hai bên; escrow milestone là tiền thanh toán cho công việc cụ thể. Hai mục đích khác nhau.'),
 ('Nếu Business không nghiệm thu thì Expert có mất tiền?', 'Hệ thống có SLA auto-approve theo cấu hình khi deliverable ở UNDER_REVIEW đủ thời gian và không có dispute/termination; Admin chạy/kiểm soát flow này.'),
 ('Tại sao không tin frontend về số dư?', 'Frontend có thể bị sửa. Backend wallet ledger mới là nguồn sự thật, hạch toán debit escrow và credit available trong transaction.'),
 ('Nếu Expert nộp trễ?', 'Backend đánh dấu OVERDUE theo deadline; chặn nộp deliverable quá hạn. Report có thể được gắn isLate để Business biết tình trạng.'),
 ('Cập nhật profile tại sao lại Pending?', 'Vì thông tin xác minh thay đổi thì phải review lại; tránh tài khoản Approved sửa sang dữ liệu không còn được xác minh.'),
 ]:
     para(d,q,lead=q); bullet(d,a)
 head(d,'12. Checklist cuối cùng: bạn có hiểu đủ để trả lời chưa?',1)
 for x in ['Nói được “vấn đề dự án giải quyết” bằng ví dụ thuê người làm việc theo milestone.', 'Phân biệt rõ account status, contract status và milestone status.', 'Giải thích được 2 lớp tiền: contract deposit và milestone escrow.', 'Nói được điều kiện để contract Active và điều kiện để Expert nhận tiền.', 'Nói được lý do reject/dispute/termination không thể chỉ đổi UI mà phải để backend quyết định.', 'Khi gặp API: nêu page → người bấm → endpoint → điều kiện backend → data/UI → lỗi/loading/quyền.']:
     bullet(d,x)
 d.core_properties.title='AITASKER - Sổ tay nghiệp vụ chi tiết từ code'
 d.core_properties.author='AITASKER Team'
 d.save(OUT); print(OUT)

if __name__=='__main__': build()
