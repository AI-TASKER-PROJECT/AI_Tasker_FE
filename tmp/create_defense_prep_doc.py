from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(r"D:\FPT\SUMMER_2026\SWP391\SO_TAY_DE_HIEU_AITASKER.docx")
NAVY, BLUE, PALE, GREY = "173F67", "2D6BA3", "F4F8FC", "5B6775"

def set_cell(cell, color=None):
    tc = cell._tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for side, val in [('top', 100), ('start', 110), ('bottom', 100), ('end', 110)]:
        x = OxmlElement(f'w:{side}'); x.set(qn('w:w'), str(val)); x.set(qn('w:type'), 'dxa'); mar.append(x)
    tc.append(mar)
    if color:
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), color); tc.append(shd)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        c=t.rows[0].cells[i]; set_cell(c, NAVY)
        r=c.paragraphs[0].add_run(h); r.bold=True; r.font.name='Arial'; r.font.size=Pt(9); r.font.color.rgb=RGBColor(255,255,255)
    for n, row in enumerate(rows):
        cells=t.add_row().cells
        for i, value in enumerate(row):
            c=cells[i]; set_cell(c, PALE if n % 2 else None)
            p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.05
            r=p.add_run(value); r.font.name='Arial'; r.font.size=Pt(8.8)
        t.rows[-1]._tr.get_or_add_trPr().append(OxmlElement('w:cantSplit'))
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

def p(doc, text, lead=None):
    para=doc.add_paragraph(); para.paragraph_format.space_after=Pt(5); para.paragraph_format.line_spacing=1.12
    if lead and text.startswith(lead):
        r=para.add_run(lead); r.bold=True; r.font.name='Arial'; r.font.size=Pt(10.5); text=text[len(lead):]
    r=para.add_run(text); r.font.name='Arial'; r.font.size=Pt(10.5)

def bullet(doc, text):
    para=doc.add_paragraph(style='List Bullet'); para.paragraph_format.space_after=Pt(3)
    r=para.add_run(text); r.font.name='Arial'; r.font.size=Pt(10.2)

def h(doc, text, level=1):
    doc.add_paragraph(text, style=f'Heading {level}')

def setup(doc):
    s=doc.sections[0]; s.top_margin=Inches(.68); s.bottom_margin=Inches(.62); s.left_margin=Inches(.72); s.right_margin=Inches(.72)
    normal=doc.styles['Normal']; normal.font.name='Arial'; normal._element.rPr.rFonts.set(qn('w:ascii'),'Arial'); normal._element.rPr.rFonts.set(qn('w:hAnsi'),'Arial'); normal.font.size=Pt(10.5)
    for name, size, color in [('Heading 1',16,NAVY),('Heading 2',12.5,BLUE),('Heading 3',11,NAVY)]:
        st=doc.styles[name]; st.font.name='Arial'; st._element.rPr.rFonts.set(qn('w:ascii'),'Arial'); st._element.rPr.rFonts.set(qn('w:hAnsi'),'Arial'); st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color); st.paragraph_format.space_before=Pt(13); st.paragraph_format.space_after=Pt(6)
    head=s.header.paragraphs[0]; head.alignment=WD_ALIGN_PARAGRAPH.RIGHT; r=head.add_run('AITASKER | SỔ TAY GIẢI THÍCH DỄ HIỂU'); r.font.name='Arial'; r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(GREY)
    foot=s.footer.paragraphs[0]; foot.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=foot.add_run('Auth • Profile • Contract • Workspace • Common'); r.font.name='Arial'; r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(GREY)

def build():
    d=Document(); setup(d)
    # cover
    q=d.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; q.paragraph_format.space_before=Pt(48); r=q.add_run('AITASKER'); r.bold=True; r.font.name='Arial'; r.font.size=Pt(14); r.font.color.rgb=RGBColor.from_string(BLUE)
    q=d.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; q.paragraph_format.space_before=Pt(20); r=q.add_run('SỔ TAY GIẢI THÍCH DỰ ÁN\nCHO NGƯỜI MỚI'); r.bold=True; r.font.name='Arial'; r.font.size=Pt(26); r.font.color.rgb=RGBColor.from_string(NAVY)
    q=d.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; q.paragraph_format.space_before=Pt(16); r=q.add_run('Phạm vi của bạn: Auth, Profile, Contract, Workspace, Common'); r.italic=True; r.font.name='Arial'; r.font.size=Pt(12); r.font.color.rgb=RGBColor.from_string(GREY)
    q=d.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; q.paragraph_format.space_before=Pt(46); r=q.add_run('Đọc theo thứ tự từ trang 2.\nMỗi phần đều trả lời: “Đây là gì? Người dùng làm gì? API là gì? Màn hình hiện gì? Lỗi thì sao?”'); r.font.name='Arial'; r.font.size=Pt(11)
    d.add_page_break()

    h(d,'1. Trước tiên: hiểu 6 từ này',1)
    p(d,'Bạn không cần thuộc code. Khi bảo vệ, chỉ cần hiểu luồng: người dùng bấm nút → trang gọi API → backend xử lý → backend trả dữ liệu → trang hiển thị kết quả.')
    add_table(d,['Từ','Nghĩa rất đơn giản','Ví dụ trong AITASKER'],[
        ['Page / màn hình','Một giao diện người dùng nhìn thấy.','Trang đăng nhập, trang hồ sơ, trang hợp đồng.'],
        ['Route','Địa chỉ để mở một page trên web.','/login là trang đăng nhập; /app/contracts là danh sách hợp đồng.'],
        ['API','“Đường dây” để giao diện hỏi backend lấy/lưu dữ liệu.','Bấm Đăng nhập → gọi POST /api/auth/login.'],
        ['Request','Thông tin frontend gửi đi.','Email và mật khẩu khi đăng nhập.'],
        ['Response','Thông tin backend trả về.','Tên người dùng, vai trò, token, trạng thái tài khoản.'],
        ['Token','Thẻ nhận diện tạm thời sau khi đăng nhập.','Frontend gửi token kèm mỗi API để backend biết ai đang gọi.'],
    ])
    h(d,'Công thức trả lời mọi câu hỏi của hội đồng',2)
    for x in ['1. Nói người dùng đang ở page nào và route nào.', '2. Nói họ bấm nút/làm thao tác gì.', '3. Nói API nào được gọi và mục đích của API đó.', '4. Nói dữ liệu trả về được đưa lên chỗ nào trên giao diện.', '5. Nói loading, lỗi, không có dữ liệu và quyền được xử lý ra sao.']:
        bullet(d,x)
    p(d,'Ví dụ ngắn: “Ở trang /login, người dùng nhập email và mật khẩu rồi bấm Đăng nhập. Trang gọi POST /api/auth/login. Nếu thành công, backend trả token, tên và role; frontend lưu session rồi đưa người dùng vào /app. Trong lúc chờ thì khóa nút; sai mật khẩu thì hiện thông báo; token hết hạn thì quay lại /login.”')

    h(d,'2. Bức tranh toàn bộ phần bạn làm',1)
    p(d,'Hãy nhớ dự án như một quá trình thuê người làm việc: đầu tiên tạo tài khoản, sau đó xác minh hồ sơ, rồi ký hợp đồng, thực hiện từng mốc công việc, cuối cùng nhận thông báo/thanh toán.')
    add_table(d,['Thứ tự','Nhóm chức năng','Ý nghĩa'],[
        ['1','Auth','Cho người dùng đăng ký, xác minh email, đăng nhập và lấy quyền truy cập.'],
        ['2','Profile','Business/Expert điền hồ sơ để hệ thống kiểm tra họ là ai và có đủ điều kiện hoạt động không.'],
        ['3','Contract','Khi Business chọn Expert, hai bên ký điều khoản, NDA và đặt cọc.'],
        ['4','Workspace','Hai bên cùng làm việc theo milestone: báo cáo, nộp sản phẩm, nghiệm thu.'],
        ['5','Common','Các phần dùng chung: menu theo quyền, thông báo, chatbot, nạp tiền.'],
    ])
    h(d,'Luồng một câu để bạn học thuộc',2)
    p(d,'“Đăng ký/đăng nhập → hoàn thiện và được duyệt hồ sơ → tạo hoặc tham gia hợp đồng → thực hiện milestone trong workspace → theo dõi thông báo và thanh toán.”')

    h(d,'3. Auth – đăng ký, đăng nhập và quên mật khẩu',1)
    p(d,'Auth giống quầy lễ tân: kiểm tra bạn là ai, tạo phiên làm việc và biết bạn thuộc vai trò nào.')
    add_table(d,['Page và route','Người dùng làm gì?','API gọi','Dữ liệu trả về hiển thị ở đâu?'],[
        ['Login\n/login','Nhập email, mật khẩu; có thể đăng nhập Google.','POST /api/auth/login\nPOST /api/auth/google/login\nGET /api/auth/me','Trả SessionUser: fullName, email, role, accountStatus, accessToken, refreshToken. Tên/role hiển thị trong AppShell; token dùng để gọi API tiếp.'],
        ['Register\n/register','Chọn Business hoặc Expert; nhập thông tin; nhận và nhập OTP.','GET /api/auth/check-email\nPOST /api/auth/email/send-otp\nPOST /api/auth/email/verify-otp\nPOST /api/auth/register','check-email kiểm tra email đã tồn tại; send-otp trả thời gian hết hạn; register trả SessionUser để đăng nhập.'],
        ['Quên mật khẩu\n/forgot-password','Nhập email để nhận hướng dẫn.','POST /api/auth/forgot-password','UI báo đã gửi email (nếu backend chấp nhận).'],
        ['Đặt lại mật khẩu\n/reset-password','Nhập mật khẩu mới theo link/token email.','POST /api/auth/reset-password','Thành công thì báo đổi mật khẩu và quay về /login.'],
    ])
    h(d,'Khi loading / lỗi / rỗng ở Auth',2)
    for x in ['Loading: nút Đăng nhập/Đăng ký bị khóa và có trạng thái đang gửi để tránh bấm nhiều lần.', 'Sai email, mật khẩu, OTP hoặc validation: hiển thị message để người dùng sửa, không xóa nội dung họ vừa nhập.', '401 Unauthorized: thường là token không hợp lệ/hết hạn. Frontend xóa session và đưa về /login.', 'Chưa đăng nhập mà vào /app/**: ProtectedRoute chặn và chuyển về /login.', 'Business/Expert chưa được duyệt hồ sơ: sau khi vào app sẽ được hướng đến phần xác minh, không được dùng đầy đủ marketplace.']:
        bullet(d,x)
    p(d,'Câu nói nên dùng: “Frontend có chặn giao diện để trải nghiệm tốt hơn, nhưng backend vẫn là nơi quyết định cuối cùng ai được phép gọi API nào.”')

    h(d,'4. Profile – hồ sơ và xác minh',1)
    p(d,'Profile là “căn cước nghề nghiệp” trên nền tảng. Business khai thông tin doanh nghiệp; Expert khai năng lực và portfolio. Staff/Admin là người duyệt.')
    add_table(d,['Page và route','API gọi','Dữ liệu dùng để hiển thị','Ai được dùng?'],[
        ['Business Profile\n/app/business/profile','GET /api/v1/profiles/business/me\nPOST /api/v1/profiles/business','Tên công ty, người đại diện, mã số thuế, giấy phép, profile/account status. Hiển thị vào form và badge trạng thái.','Business'],
        ['Business Verification\n/app/business/kyb','POST /api/v1/profiles/business/license-file\nGET /api/auth/tax-check/:taxCode','URL/tệp giấy phép để preview; kết quả kiểm tra mã số thuế.','Business'],
        ['Expert Profile\n/app/expert/profile','GET /api/v1/profiles/expert/me\nPOST /api/v1/profiles/expert','Chức danh, giới thiệu, kỹ năng, lĩnh vực/công nghệ, trạng thái profile.','Expert'],
        ['Expert Portfolio\n/app/expert/portfolio','GET /api/v1/profiles/portfolio/me\nPOST /api/v1/profiles/portfolio\nPOST /api/v1/profiles/portfolio/certificate-file','Kinh nghiệm, dự án, kỹ năng và URL chứng chỉ. Hiển thị thành portfolio.','Expert'],
        ['Duyệt hồ sơ\n/app/verifications\n/app/verifications/:type/:id','GET /api/v1/profiles/business\nGET /api/v1/profiles/expert\nGET /api/v1/profiles/files/view-url\nPOST /api/v1/profiles/approve/:type/:profileId?status=...','Danh sách hồ sơ, thông tin chi tiết và link xem file. Sau approve/reject, UI cập nhật badge trạng thái/lý do từ chối.','Staff, Admin'],
    ])
    h(d,'Nếu dữ liệu chưa có hoặc có lỗi',2)
    for x in ['Chưa có profile: hiển thị form trống hoặc lời nhắc tạo hồ sơ, không giả lập dữ liệu nghiệp vụ.', 'Đang lấy profile: hiển thị loading trước khi đổ dữ liệu vào form.', 'Upload file lỗi: báo lỗi file/tải lên; giữ các trường form để người dùng không phải nhập lại.', 'Reject hồ sơ: Staff/Admin phải nhập reason. Business/Expert xem được lý do để chỉnh sửa và gửi lại.', 'Approved có nghĩa là hồ sơ đã được xác minh. Đây là điều kiện để Business/Expert đi sâu vào nghiệp vụ giao dịch.']:
        bullet(d,x)

    h(d,'5. Contract – hợp đồng giữa Business và Expert',1)
    p(d,'Contract là phần “chốt thỏa thuận”. Nó không phải nơi làm việc hằng ngày; nó lưu điều khoản, hai bên ký, NDA, đặt cọc và tổng quan các milestone.')
    add_table(d,['Page và route','API gọi','Dữ liệu trả về hiển thị gì?','Lỗi/loading/quyền'],[
        ['Danh sách hợp đồng\n/app/contracts','GET /api/v1/contracts','Mỗi Contract có contractId, jobId, Business, Expert, ngân sách, thời gian, status, progress, milestones. UI hiển thị danh sách/lọc theo status.','Business/Expert xem hợp đồng liên quan. Rỗng: “Chưa có hợp đồng”.'],
        ['Chi tiết hợp đồng\n/app/contracts/:contractId','GET /api/v1/contracts/:id\nGET /api/v1/contracts/:id/milestones','Điều khoản, ngân sách, bên tham gia, trạng thái ký/NDA/đặt cọc, danh sách mốc.','Đang tải: loading. Không tìm thấy ID: thông báo lỗi. Nút hành động chỉ mở cho role/status phù hợp.'],
        ['Ký / NDA / từ chối','POST /api/v1/contracts/:id/sign\nPOST /api/v1/contracts/:id/nda-sign\nPOST /api/v1/contracts/:id/reject\nPOST /api/v1/contracts/:id/cancel-draft','Response trả contract với status mới/timestamp ký. UI tải lại detail và chỉ hiện bước tiếp theo hợp lệ.','Khóa nút khi đang gửi. Backend báo status không hợp lệ: hiện message, UI không tự đổi trạng thái.'],
        ['Đặt cọc','POST /api/v1/contracts/:id/deposit/pay\nPOST /api/v1/contracts/:id/expert-deposit/pay','PaymentActionResponse: completed, needTopup, currentBalance, requiredAmount, missingAmount, redirectUrl. UI biết đã trả hay phải nạp ví.','Nếu needTopup=true hoặc không đủ tiền: mở/hướng dẫn top-up, không giả vờ thanh toán thành công.'],
    ])
    p(d,'Câu trả lời dễ nhớ: “Contract Detail quản lý thỏa thuận trước và trong khi bắt đầu; Workspace quản lý công việc thực tế theo từng milestone.”')

    h(d,'6. Workspace – nơi hai bên thực sự làm việc',1)
    p(d,'Route chính: /app/contracts/:contractId/workspace. Milestone hiểu đơn giản là một mốc công việc: có hạn, ngân sách, tiêu chí nghiệm thu và sản phẩm cần nộp.')
    add_table(d,['Việc trên Workspace','API quan trọng','UI hiển thị gì?','Ai làm?'],[
        ['Xem mốc và tiêu chí','GET /api/v1/contracts/:contractId/milestones\nGET /api/v1/milestones/:milestoneId/criteria\nGET /api/v1/milestones/:milestoneId/deliverables','Timeline, deadline, ngân sách, status, acceptance criteria, sản phẩm đã nộp.','Business và Expert xem'],
        ['Bắt đầu mốc / ký quỹ','POST .../milestones/:milestoneId/deposit\nPOST /api/v1/milestones/:milestoneId/start','Status milestone mới và thông tin ký quỹ. UI chuyển trạng thái khi backend xác nhận.','Business deposit; Expert start'],
        ['Báo cáo tiến độ','POST .../progress-report-request\nPOST/GET .../progress-reports\nPOST .../:reportId/acknowledge\nPOST .../feedback','Nội dung report, thời gian gửi, đã xác nhận chưa, feedback. UI là lịch sử báo cáo/badge chờ phản hồi.','Expert gửi; Business yêu cầu/xác nhận/feedback'],
        ['Nộp sản phẩm','POST .../source-code-file\nPOST .../deliverables\nPOST /api/v1/milestones/:milestoneId/complete','Link/tệp source code hoặc deliverable, thời điểm nộp, trạng thái chờ review.','Expert'],
        ['Nghiệm thu / từ chối','POST .../milestones/:milestoneId/approve\nPOST .../milestones/:milestoneId/reject','Status mới; nếu reject có reason và failedCriteria. UI hiển thị kết quả/lý do.','Business'],
        ['Đổi yêu cầu / chấm dứt','GET/POST .../change-requests\nPOST .../termination-requests','ChangeRequest/TerminationRequest: reason, status, review note, settlement. UI dùng form/modal và lịch sử action.','Business/Expert; Staff/Admin theo quyền backend'],
    ])
    h(d,'Một tình huống để giải thích khi bị hỏi',2)
    p(d,'“Expert nộp deliverable. Business mở Workspace, xem file và các tiêu chí. Nếu đạt, Business gọi API approve và milestone chuyển sang hoàn thành. Nếu chưa đạt, Business gọi reject kèm lý do; UI hiển thị lý do để Expert sửa và nộp lại. Vì trạng thái ảnh hưởng tiền và quyền, UI luôn chờ backend trả kết quả rồi mới cập nhật.”')

    h(d,'7. Common – các chức năng dùng chung',1)
    add_table(d,['Vị trí','API','Dữ liệu/UI','Điểm quan trọng'],[
        ['AppShell: sidebar, user menu','Dùng SessionUser hiện tại; profile refresh khi cần','Tên, role, accountStatus; menu đổi theo quyền.','Business/Expert chưa Approved được dẫn về verification.'],
        ['Thông báo\n/app/notifications','GET /api/v1/notifications\nGET /api/v1/notifications/unread-count\nPATCH /api/v1/notifications/:id/read\nPATCH /api/v1/notifications/read-all','NotificationItem[] cho dropdown/list; unread count là badge số chưa đọc.','Rỗng: thông báo chưa có thông báo; socket/reload giúp cập nhật.'],
        ['Chatbox','POST /api/chatbot/ask','answer và sources được thêm vào khung chat.','Khóa nút gửi khi chờ; lỗi hiện fallback message và giữ lại câu hỏi.'],
        ['Nạp tiền ví','POST /api/payments/payos/create\nPOST /api/payments/payos/:orderCode/sync\nGET /api/v1/wallet/me','Checkout URL/QR, orderCode và số dư ví. UI hiện modal thanh toán rồi tải lại số dư sau sync.','Frontend không tự cộng tiền; số dư chỉ tin backend sau khi xác nhận thanh toán.'],
    ])

    h(d,'8. Quyền: Business, Expert, Staff, Admin khác nhau thế nào?',1)
    p(d,'Nhớ theo câu hỏi “ai là người thuê, ai là người làm, ai là người kiểm tra, ai là người quản trị?”.')
    add_table(d,['Role','Nói đơn giản','Việc chính trong phần bạn phụ trách'],[
        ['Business','Bên thuê chuyên gia.','Tạo/hoàn thiện hồ sơ doanh nghiệp; ký hợp đồng/NDA; đặt cọc; yêu cầu và phản hồi report; duyệt hoặc từ chối milestone.'],
        ['Expert','Bên nhận việc.','Hoàn thiện hồ sơ/portfolio; ký hợp đồng/NDA; bắt đầu milestone; gửi report; nộp source code/deliverable; phản hồi change/termination.'],
        ['Staff','Nhân sự vận hành.','Duyệt hoặc từ chối hồ sơ xác minh; xử lý các case được giao như tranh chấp theo quy trình.'],
        ['Admin','Quản trị hệ thống.','Có quyền quản trị rộng hơn: duyệt hồ sơ, quản lý account/staff/setting/master data/wallet/reports tùy endpoint backend.'],
    ])
    p(d,'Điểm quan trọng: menu frontend có thể ẩn chức năng không đúng role, nhưng backend vẫn kiểm tra token và role. Vì vậy người dùng tự sửa URL cũng không thể làm hành động nếu backend không cho phép.')

    h(d,'9. 6 câu phản biện hay gặp – trả lời ngắn gọn',1)
    for question, answer in [
        ('API lỗi thì hệ thống làm gì?', 'Hiện message dễ hiểu, dừng loading, giữ lại dữ liệu form để người dùng sửa/gửi lại. Với 401 thì xóa session và login lại; 403 báo không đủ quyền.'),
        ('Dữ liệu rỗng thì sao?', 'Không hiển thị bảng trống khó hiểu. UI có empty state và nút hành động, ví dụ “Tạo hồ sơ” hoặc “Chưa có hợp đồng”.'),
        ('Tại sao không đổi trạng thái ngay trên frontend?', 'Vì contract, payment và milestone liên quan tiền/quyền. Backend là nguồn đúng duy nhất; UI chờ response rồi cập nhật.'),
        ('Tại sao cần Approved profile?', 'Để hạn chế tài khoản chưa xác minh tham gia giao dịch, tạo niềm tin và giảm rủi ro tranh chấp.'),
        ('Khác nhau giữa Contract và Workspace?', 'Contract là thỏa thuận: điều khoản, chữ ký, NDA, đặt cọc. Workspace là thực thi: milestone, report, nộp sản phẩm, nghiệm thu.'),
        ('Bảo mật ở đâu?', 'Token được gắn vào request. Frontend bảo vệ route cho trải nghiệm; backend xác thực JWT và phân quyền cho mọi API quan trọng.'),
    ]:
        p(d, question, lead=question); bullet(d, answer)
    h(d,'10. Checklist 5 phút trước khi bảo vệ',1)
    for x in ['Chuẩn bị sẵn 1 tài khoản Business và 1 Expert; nếu demo duyệt hồ sơ thì có Staff/Admin.', 'Chuẩn bị dữ liệu: profile Approved, một contract có milestone, một deliverable/report nếu có.', 'Mở theo luồng: login → profile → contracts → contract detail → workspace → notifications/top-up.', 'Khi được hỏi một page, dùng đúng công thức 5 bước ở trang 2.', 'Nếu dịch vụ email/PayOS không ổn định, chuẩn bị ảnh/video hoặc dữ liệu mẫu làm phương án dự phòng.']:
        bullet(d,x)
    d.core_properties.title='AITASKER - Sổ tay giải thích dự án cho người mới'
    d.core_properties.author='AITASKER Team'
    d.save(OUT)
    print(OUT)

if __name__ == '__main__':
    build()
