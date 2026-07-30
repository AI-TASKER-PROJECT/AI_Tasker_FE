from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from create_business_handbook import setup, table, para, bullet, head, box

OUT=Path(r"D:\FPT\SUMMER_2026\SWP391\API_PAGE_OWNERSHIP_AITASKER.docx")

def title(d, text, route, file_path):
    head(d,text,1); para(d,f'Nguồn service: {file_path}.',lead='Nguồn service: '); para(d,f'Các page/route liên quan: {route}.',lead='Các page/route liên quan: ')

def build():
 d=Document(); setup(d)
 q=d.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; q.paragraph_format.space_before=Pt(48); r=q.add_run('AITASKER'); r.bold=True; r.font.name='Arial'; r.font.size=Pt(14); r.font.color.rgb=RGBColor.from_string('2D6BA3')
 q=d.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; q.paragraph_format.space_before=Pt(18); r=q.add_run('BẢN ĐỒ API – PAGE\nĐANG SỞ HỮU'); r.bold=True; r.font.name='Arial'; r.font.size=Pt(25); r.font.color.rgb=RGBColor.from_string('173F67')
 q=d.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; q.paragraph_format.space_before=Pt(16); r=q.add_run('Auth • Profile • Contract • Workspace • Common'); r.italic=True; r.font.name='Arial'; r.font.size=Pt(12)
 q=d.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; q.paragraph_format.space_before=Pt(38); r=q.add_run('Mỗi dòng trả lời: API ở service nào, page nào đang gọi, và API đó phục vụ việc gì.'); r.font.name='Arial'; r.font.size=Pt(10.5)
 d.add_page_break()

 head(d,'1. Quy ước đọc tài liệu',1)
 table(d,['Cột','Ý nghĩa'],[
 ['Source service','File TypeScript khai báo endpoint. Đây là nơi gom API, không phải UI.'],
 ['Page/component sở hữu','Màn hình hoặc component gọi API để xử lý giao diện. Một API có thể được dùng ở nhiều page.'],
 ['Không có page chính','API có thể phục vụ flow Marketplace, Risk hoặc Admin; không gán sai sang page bạn phụ trách.'],
 ['Response','Frontend thường nhận riêng `data` vì `apiClient.call<T>()` bỏ wrapper `{ success, message, data }`.'],
 ])
 box(d,'Lưu ý khi bảo vệ:', 'Contract và Workspace cùng dùng `contractService.ts`. Contract Detail chịu trách nhiệm thỏa thuận/ký/ký quỹ/change request; Workspace chịu trách nhiệm milestone, report, deliverable, nghiệm thu và case đang chạy.')

 title(d,'2. Auth','/login • /register • /forgot-password • /reset-password • AppShell/session','src/services/authServices.ts')
 table(d,['API','Hàm service','Page/component gọi','Chức năng và data chính'],[
 ['POST /api/auth/login','authApi.login','LoginPage (/login)','Gửi { email, password }; nhận SessionUser { accessToken, refreshToken, role, accountStatus, email, fullName }; lưu session và điều hướng.'],
 ['GET /api/auth/check-email?email=…','authApi.checkEmail','LoginPage, RegisterPage, ForgotPasswordPage','Nhận boolean email tồn tại; Login/forgot kiểm tra account, Register chặn email trùng.'],
 ['POST /api/auth/google/login','authApi.googleLogin','LoginPage, RegisterPage','Gửi Google credential và role tùy luồng; nhận SessionUser.'],
 ['POST /api/auth/google/register','authApi.googleSignup','LoginPage, RegisterPage','Tạo account Google mới với credential/fullName/phone/role.'],
 ['POST /api/auth/email/send-otp','authApi.sendOtp','RegisterPage','Gửi { email }; nhận { expiresIn } để UI mở bước OTP/đếm hạn.'],
 ['POST /api/auth/email/verify-otp','authApi.verifyOtp','RegisterPage','Gửi { email, otp }; xác nhận email trước khi register.'],
 ['POST /api/auth/register','authApi.register','RegisterPage','Gửi { email,password,fullName,phone,role }; trả SessionUser, account mới Pending.'],
 ['POST /api/auth/forgot-password','authApi.forgotPassword','ForgotPasswordPage','Gửi { email }; backend gửi hướng dẫn reset.'],
 ['POST /api/auth/reset-password','authApi.resetPassword','ResetPasswordPage','Gửi { token, newPassword }; thành công đưa user về login.'],
 ['GET /api/auth/me','authApi.me','AppShell/session flow','Đồng bộ role, accountStatus, email, fullName của session hiện tại.'],
 ['POST /api/auth/refresh','authApi.refresh','Session/token refresh flow','Gửi refreshToken để nhận access token/session mới.'],
 ])

 title(d,'3. Profile và Verification','/app/business/profile • /app/business/kyb • /app/expert/profile • /app/expert/portfolio • /app/verifications','src/services/profileService.ts')
 table(d,['API','Hàm service','Page/component gọi','Chức năng'],[
 ['GET /api/v1/profiles/business/me','getMyBusiness','BusinessProfilePage, BusinessVerificationProfilePage','Đổ BusinessProfile vào form, badge KYB/status, giấy phép.'],
 ['POST /api/v1/profiles/business','upsertBusiness','BusinessProfilePage, BusinessVerificationProfilePage','Lưu KYB. Gửi taxCode, businessLicenseUrl và profile; trả BusinessProfile mới.'],
 ['POST /api/v1/profiles/business/license-file','uploadBusinessLicense','BusinessProfilePage, BusinessVerificationProfilePage','Upload FormData { file }; trả URL để gắn vào businessLicenseUrl.'],
 ['GET /api/auth/tax-check/:taxCode','checkTaxCode','BusinessVerificationProfilePage','Trả companyName/address/representative để preview khi nhập mã số thuế.'],
 ['GET /api/v1/profiles/expert/me','getMyExpert','ExpertProfilePage, ExpertPortfolioPage','Đổ ExpertProfile/KYC vào form.'],
 ['POST /api/v1/profiles/expert','upsertExpert','ExpertProfilePage, ExpertVerificationProfilePage','Lưu nationalId, portfolioUrl, yearsOfExperience; trả ExpertProfile/status.'],
 ['POST /api/v1/profiles/expert/portfolio-file','uploadExpertPortfolio','ExpertProfilePage','Upload file năng lực, trả URL portfolioUrl.'],
 ['GET /api/v1/profiles/portfolio/me','getMyPortfolio','ExpertProfilePage, ExpertPortfolioPage','Đổ portfolio: domain/skill/technology/certificates/selfDescription.'],
 ['POST /api/v1/profiles/portfolio','upsertPortfolio','ExpertPortfolioPage','Lưu năng lực, domainIds, skillIds, technologyIds, certificates.'],
 ['POST /api/v1/profiles/portfolio/certificate-file','uploadExpertCertificate','ExpertPortfolioPage','Upload chứng chỉ, trả URL.'],
 ['GET /api/v1/profiles/business','listBusinesses','VerificationsPage','Lấy danh sách Business để Staff lọc và mở chi tiết review.'],
 ['GET /api/v1/profiles/expert','listExperts','VerificationsPage','Lấy danh sách Expert để Staff lọc và mở chi tiết review.'],
 ['GET /api/v1/profiles/business/:id','getBusinessById','VerificationDetailPage, ContractDetailPage, public profile flow','Lấy Business detail; Contract Detail dùng để hiện đối tác.'],
 ['GET /api/v1/profiles/expert/:id','getExpertById','VerificationDetailPage, ContractDetailPage, public profile flow','Lấy Expert detail; Contract Detail dùng để hiện đối tác.'],
 ['GET /api/v1/profiles/business/by-job/:jobId','getBusinessByJob','Job Detail/Submit Proposal flow','Expert xem thông tin Business sở hữu job trước khi proposal.'],
 ['GET /api/v1/profiles/portfolio','listPortfolios','VerificationDetailPage, public profile flow','Tìm portfolio theo expertId để hiển thị năng lực.'],
 ['GET /api/v1/profiles/files/view-url?path=…','getFileViewUrl','VerificationDetailPage','Lấy URL xem giấy phép/portfolio/chứng chỉ an toàn.'],
 ['POST /api/v1/profiles/approve/:type/:profileId','approve','VerificationDetailPage','Staff approve/reject với query status/reason; cập nhật badge và rejection reason.'],
 ['POST /api/v1/proposals/file','uploadProposalFile','SubmitProposalPage','Upload file đính kèm proposal; thuộc Profile service nhưng được Marketplace dùng.'],
 ])
 para(d,'Route frontend cho `/app/verifications` cho Staff/Admin, nhưng code backend của thao tác approve hiện yêu cầu STAFF. Khi nói bảo vệ nên nêu đúng hai lớp này.')

 title(d,'4. Contract – Contract Detail và danh sách hợp đồng','/app/contracts • /app/contracts/:contractId • ProposalDetailPage • /app/reviews','src/services/contractService.ts')
 table(d,['API','Hàm service','Page/component gọi','Chức năng'],[
 ['GET /api/v1/contracts','listContracts','ContractsPage; Contract Detail/Workspace refresh','Danh sách contract, dùng status/progress/budget/party để tạo card và filter.'],
 ['GET /api/v1/contracts/:contractId','getContract','ContractDetailPage, WorkspacePage, WalletPage enrichment','Lấy điều khoản, party, chữ ký/NDA, status, tổng ngân sách và progress.'],
 ['POST /api/v1/contracts/from-proposals/:proposalId','createFromProposal','ProposalDetailPage / Business chọn proposal','Tạo DRAFT contract từ proposal Accepted.'],
 ['POST /api/v1/contracts/:id/sign','sign','ContractDetailPage','Business/Expert xác nhận chữ ký contract.'],
 ['POST /api/v1/contracts/:id/nda-sign','signNda','ContractDetailPage','Business/Expert ký NDA; đủ chữ ký/NDA thì chuyển PENDING.'],
 ['POST /api/v1/contracts/:id/reject','rejectContract','ContractDetailPage','Expert từ chối contract.'],
 ['POST /api/v1/contracts/:id/cancel-draft','cancelDraft','ContractDetailPage','Business hủy nháp hợp lệ.'],
 ['POST /api/v1/contracts/:id/deposit/pay','payDeposit','ContractDetailPage','Business ký quỹ 20% contract; trả PaymentActionResponse để xử lý top-up.'],
 ['POST /api/v1/contracts/:id/expert-deposit/pay','payExpertDeposit','ContractDetailPage','Expert ký quỹ 10% contract.'],
 ['POST /api/v1/admin/contracts/:id/deposits/refund','refundContractDeposits','Admin/Risk flow, không phải Contract Detail chính','Hoàn ký quỹ do Admin xử lý.'],
 ['POST /api/v1/contracts/:id/change-requests','createChangeRequest','ContractDetailPage','Tạo yêu cầu thay đổi scope/budget/timeline/milestone.'],
 ['GET /api/v1/contracts/:id/change-requests','listChangeRequests','ContractsPage, ContractDetailPage','Lấy lịch sử thay đổi để hiện badge/list.'],
 ['POST .../change-requests/:requestId/accept','acceptChangeRequest','ContractDetailPage','Chấp nhận change request, body { reviewNote }; page reload contract/milestone.'],
 ['POST .../change-requests/:requestId/reject','rejectChangeRequest','ContractDetailPage','Từ chối change request, body { reviewNote }.'],
 ['POST/GET /api/v1/contracts/:id/reviews','createReview/listReviews','ReviewsPage','Tạo/lấy review sau hợp đồng.'],
 ])

 title(d,'5. Workspace – milestone, report, deliverable và nghiệm thu','/app/contracts/:contractId/workspace','src/services/contractService.ts')
 table(d,['API','Hàm service','Page/component gọi','Chức năng'],[
 ['GET /api/v1/contracts/:id/milestones','listMilestones','WorkspacePage, ContractDetailPage, ContractsPage','Lấy milestone đã chốt, status, finalBudget, criteria snapshot, reject/resubmit count.'],
 ['GET /api/v1/jobs/:jobId/milestones','listJobMilestones','ContractDetailPage, CreateJob/ManageJob flow','Lấy milestone gốc của job để đối chiếu.'],
 ['GET /api/v1/milestones/:id/criteria','listCriteria','WorkspacePage','Checklist acceptance criteria cho mốc được chọn.'],
 ['POST/PUT/DELETE criteria','createCriteria/updateCriteria/deleteCriteria','CreateJobPage/ManageJob flow','Tạo/sửa/xóa tiêu chí nghiệm thu của job milestone.'],
 ['POST /api/v1/contracts/:c/milestones/:m/deposit','depositMilestoneEscrow','WorkspacePage','Business giữ finalBudget của milestone vào escrow, cập nhật status.'],
 ['POST /api/v1/milestones/:id/start','startMilestone','WorkspacePage','Expert bắt đầu milestone khi điều kiện hợp lệ.'],
 ['POST .../progress-report-request','requestProgressReport','WorkspacePage','Business yêu cầu report; trả dueAt/request status.'],
 ['POST .../progress-reports','submitProgressReport','WorkspacePage','Expert gửi content, %, attachment/source/demo/notes; trả report và acknowledgementState.'],
 ['GET .../progress-reports','listProgressReports','WorkspacePage','Hiện lịch sử tiến độ, late/feedback/acknowledgement.'],
 ['POST .../progress-reports/:reportId/acknowledge','acknowledgeProgressReport','WorkspacePage','Business xác nhận report đã xem.'],
 ['POST .../progress-reports/:reportId/feedback','feedbackProgressReport','WorkspacePage','Business gửi feedback/category/severity/DoD/requiresAdjustment.'],
 ['POST .../source-code-file','uploadMilestoneSourceCode','WorkspacePage','Upload ZIP source bằng FormData; trả URL.'],
 ['POST .../deliverables','submitDeliverable','WorkspacePage','Expert nộp deliverable với source/demo/notes; trả submissionRound/status.'],
 ['GET /api/v1/milestones/:id/deliverables','listDeliverables','WorkspacePage','Hiện các lần nộp, status và rejectionFeedback.'],
 ['POST .../approve','approveMilestone','WorkspacePage','Business nghiệm thu; backend release escrow 100% sang Expert.'],
 ['POST .../reject','rejectMilestone','WorkspacePage','Business từ chối với reason/failedCriteria; mốc quay về IN_PROGRESS.'],
 ['PATCH /api/v1/jobs/:jobId/milestones/:id','updateMilestone','CreateJobPage/ManageJobPage','Sửa milestone của job khi chưa bị khóa bởi contract.'],
 ['POST .../milestones/check-overdue','checkOverdueMilestones','WorkspacePage khi Admin thao tác','Đánh dấu mốc quá hạn theo deadline.'],
 ['POST .../milestones/sla-auto-approve','autoApproveReviewSla','WorkspacePage khi Admin thao tác','Tự nghiệm thu mốc quá SLA nếu không có case chặn.'],
 ])

 title(d,'6. Workspace – chấm dứt và bằng chứng','/app/contracts/:contractId/workspace • Risk/Admin pages','src/services/contractService.ts')
 table(d,['API','Hàm service','Page/component gọi','Chức năng'],[
 ['POST /contracts/:id/termination-requests','requestTermination','WorkspacePage','Business/Expert tạo yêu cầu chấm dứt.'],
 ['POST /contracts/:id/immediate-termination','immediateTermination','WorkspacePage','Chấm dứt ngay, body { reason, confirmedPenalty }.'],
 ['GET /contracts/:id/termination-requests','listTerminationRequests','WorkspacePage','Hiện lịch sử và status yêu cầu chấm dứt.'],
 ['GET /termination-requests/:id','getTerminationRequest','Workspace/Risk detail flow','Lấy detail một request.'],
 ['POST .../:id/accept hoặc /dispute','acceptTerminationRequest/disputeTerminationRequest','WorkspacePage','Bên còn lại đồng ý hoặc phản đối kèm reason.'],
 ['POST .../:id/assign-staff, /reject, /approve','assign/reject/approveTerminationRequest','Risk/Admin/Staff flow','Giao Staff, từ chối hoặc duyệt case tùy quyền backend.'],
 ['POST .../:id/partial-evidence','submitTerminationPartialEvidence','Risk/termination flow','Nộp bằng chứng phần công việc đã làm.'],
 ['POST .../:id/execute-settlement','executeTerminationSettlement','Admin/Risk flow','Thực thi chia tiền/refund khi termination.'],
 ['POST .../:id/withdraw hoặc /refund-deposit','withdraw/refundTerminationDeposit','Workspace/Admin flow','Rút yêu cầu hoặc hoàn ký quỹ sau termination.'],
 ['GET/POST /api/v1/case-attachments','listCaseAttachments/createCaseAttachment','DisputeDetailPage/Risk flow','Lấy/thêm evidence cho dispute hoặc termination.'],
 ['POST /termination-requests/expire-awaiting-expert','expireAwaitingExpertTerminationRequests','Admin scheduled/operational flow','Xử lý request chờ Expert phản hồi đã quá hạn.'],
 ])

 title(d,'7. Common – thông báo, ví, PayOS, chatbot và client chung','/app/notifications • /app/wallet • AppShell top-up/chatbox • toàn app','notificationService.ts • walletService.ts • paymentService.ts • chatbotService.ts • apiClient.ts')
 table(d,['API / file','Hàm service','Page/component gọi','Chức năng'],[
 ['GET /api/v1/notifications','notificationApi.list','NotificationsPage, AppShell','Lấy NotificationItem[]: type/title/message/targetUrl/isRead/time.'],
 ['GET /api/v1/notifications/unread-count','notificationApi.unreadCount','AppShell badge','Lấy số thông báo chưa đọc.'],
 ['PATCH /notifications/:id/read','notificationApi.markRead','NotificationsPage','Đánh dấu 1 notification đã đọc.'],
 ['PATCH /notifications/read-all','notificationApi.markAllRead','NotificationsPage','Đánh dấu toàn bộ đã đọc.'],
 ['GET /api/v1/wallet/me','walletApi.current','WalletPage, ContractDetailPage, AppShell top-up refresh','Lấy current/available/escrow/holding/disputed balance.'],
 ['POST /api/payments/payos/create','paymentApi.createWalletTopup','Top-up modal trong AppShell/Contract flow','Gửi { amount, description }, nhận checkout URL/QR/orderCode.'],
 ['POST /api/payments/payos/:orderCode/sync','paymentApi.syncWalletTopup','Top-up modal','Đồng bộ PaymentOrder; chỉ sau đó reload wallet.'],
 ['POST /api/chatbot/ask','chatbotApi.ask','AppShell Chatbox','Gửi { question }, nhận { answer, sources }.'],
 ['Toàn bộ endpoint','apiClient.call / interceptor','Toàn bộ page','Gắn Bearer token, unwrap `data`, map lỗi, 401 clearSession và redirect login.'],
 ])
 head(d,'8. Kết luận ngắn để nói khi bảo vệ',1)
 for x in ['Service file là nơi khai báo API dùng lại; page là nơi điều khiển lúc nào gọi, loading/lỗi/rỗng và đưa response lên UI.', 'ContractDetailPage sở hữu thỏa thuận: ký, NDA, ký quỹ hợp đồng, change request. WorkspacePage sở hữu thực thi: escrow milestone, report, deliverable, approve/reject.', 'API trả data đã được apiClient bỏ wrapper. Page không tự “đổi tiền” hoặc status quan trọng; sau action luôn lấy response/reload từ backend.']:
   bullet(d,x)
 d.core_properties.title='AITASKER - Bản đồ API và Page đang sở hữu'
 d.core_properties.author='AITASKER Team'
 d.save(OUT); print(OUT)

if __name__=='__main__': build()
