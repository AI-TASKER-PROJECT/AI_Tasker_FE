from pathlib import Path
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from create_business_handbook import setup, table, para, head, box

ROOT = Path(r"D:\FPT\SUMMER_2026\SWP391\Code\AI_Tasker_FE")
OUT = Path(r"D:\FPT\SUMMER_2026\SWP391\API_CATALOG_TOAN_BO_AUTH_PROFILE_CONTRACT_WORKSPACE_COMMON.docx")

SOURCES = [
    ("AUTH", "src/services/authServices.ts", "authApi"),
    ("PROFILE", "src/services/profileService.ts", "profileApi"),
    ("CONTRACT + WORKSPACE", "src/services/contractService.ts", "contractApi"),
    ("COMMON - Notification", "src/services/notificationService.ts", "notificationApi"),
    ("COMMON - Chatbot", "src/services/chatbotService.ts", "chatbotApi"),
]

MEANING = {
 "login":"Đăng nhập email/mật khẩu", "me":"Lấy session người đang đăng nhập", "refresh":"Làm mới access token", "checkEmail":"Kiểm tra email đã tồn tại", "sendOtp":"Gửi OTP email", "verifyOtp":"Xác minh OTP", "forgotPassword":"Gửi yêu cầu quên mật khẩu", "resetPassword":"Đặt lại mật khẩu", "register":"Đăng ký tài khoản", "googleSignup":"Đăng ký bằng Google", "googleLogin":"Đăng nhập Google",
 "upsertBusiness":"Tạo hoặc cập nhật hồ sơ Business (KYB)", "getMyBusiness":"Lấy KYB của Business hiện tại", "uploadBusinessLicense":"Tải giấy phép kinh doanh", "upsertExpert":"Tạo hoặc cập nhật hồ sơ Expert (KYC)", "getMyExpert":"Lấy KYC của Expert hiện tại", "uploadExpertPortfolio":"Tải file portfolio Expert", "upsertPortfolio":"Tạo hoặc cập nhật portfolio", "getMyPortfolio":"Lấy portfolio hiện tại", "uploadExpertCertificate":"Tải chứng chỉ", "uploadProposalFile":"Tải file đính kèm proposal", "listBusinesses":"Danh sách Business", "getBusinessByJob":"Lấy Business theo Job", "getBusinessById":"Lấy Business theo ID", "listBusinessJobs":"Lọc Job của Business ở phía frontend", "listExperts":"Danh sách Expert", "getExpertById":"Lấy Expert theo ID", "listPortfolios":"Danh sách portfolio", "getPortfolioByExpert":"Tìm portfolio Expert trong danh sách", "getFileViewUrl":"Lấy URL xem file", "approve":"Duyệt hoặc từ chối profile", "checkTaxCode":"Tra mã số thuế",
 "listContracts":"Danh sách hợp đồng", "getContract":"Chi tiết hợp đồng", "createFromProposal":"Tạo hợp đồng từ proposal", "sign":"Ký hợp đồng", "signNda":"Ký NDA", "rejectContract":"Từ chối ký hợp đồng", "cancelDraft":"Hủy hợp đồng nháp", "payDeposit":"Business thanh toán đặt cọc", "payExpertDeposit":"Expert thanh toán đặt cọc", "refundContractDeposits":"Admin hoàn tiền đặt cọc hợp đồng", "immediateTermination":"Yêu cầu chấm dứt ngay", "requestTermination":"Tạo yêu cầu chấm dứt", "listTerminationRequests":"Danh sách yêu cầu chấm dứt", "getTerminationRequest":"Chi tiết yêu cầu chấm dứt", "disputeTerminationRequest":"Tranh chấp yêu cầu chấm dứt", "acceptTerminationRequest":"Chấp nhận yêu cầu chấm dứt", "assignTerminationStaff":"Giao Staff xử lý chấm dứt", "rejectTerminationRequest":"Từ chối yêu cầu chấm dứt", "approveTerminationRequest":"Admin phê duyệt chấm dứt", "submitTerminationPartialEvidence":"Nộp bằng chứng một phần", "executeTerminationSettlement":"Thực hiện settlement chấm dứt", "withdrawTerminationRequest":"Rút yêu cầu chấm dứt", "refundTerminationDeposit":"Hoàn đặt cọc khi chấm dứt", "listCaseAttachments":"Danh sách file hồ sơ vụ việc", "createCaseAttachment":"Tạo file hồ sơ vụ việc", "expireAwaitingExpertTerminationRequests":"Hết hạn yêu cầu chờ Expert", "createMilestone":"Tạo milestone độc lập", "createJobMilestone":"Tạo milestone cho Job", "listMilestones":"Danh sách milestone của contract", "listJobMilestones":"Danh sách milestone của Job", "createCriteria":"Tạo tiêu chí nghiệm thu", "updateCriteria":"Sửa tiêu chí nghiệm thu", "deleteCriteria":"Xóa tiêu chí nghiệm thu", "listCriteria":"Danh sách tiêu chí nghiệm thu", "submitDeliverable":"Nộp deliverable", "uploadMilestoneSourceCode":"Tải file source code", "submitProgressReport":"Expert nộp báo cáo tiến độ", "requestProgressReport":"Business yêu cầu báo cáo", "listProgressReports":"Danh sách báo cáo tiến độ", "acknowledgeProgressReport":"Business xác nhận đã xem báo cáo", "feedbackProgressReport":"Business phản hồi báo cáo", "listDeliverables":"Danh sách deliverable", "depositMilestoneEscrow":"Ký quỹ milestone", "startMilestone":"Bắt đầu milestone", "approveMilestone":"Business duyệt milestone", "rejectMilestone":"Business từ chối milestone", "completeMilestone":"Hoàn tất milestone", "updateMilestone":"Sửa milestone Job", "createChangeRequest":"Tạo yêu cầu thay đổi contract", "listChangeRequests":"Danh sách yêu cầu thay đổi", "acceptChangeRequest":"Chấp nhận yêu cầu thay đổi", "rejectChangeRequest":"Từ chối yêu cầu thay đổi", "checkOverdueMilestones":"Kiểm tra milestone quá hạn", "autoApproveReviewSla":"Tự duyệt theo SLA", "createReview":"Tạo review hợp đồng", "listReviews":"Danh sách review hợp đồng",
 "list":"Danh sách notification", "unreadCount":"Đếm notification chưa đọc", "markRead":"Đánh dấu một notification đã đọc", "markAllRead":"Đánh dấu tất cả notification đã đọc", "ask":"Gửi câu hỏi tới chatbot",
}

def compact(s):
    return re.sub(r"\s+", " ", s).strip()

def parse_methods(path):
    text = path.read_text(encoding="utf-8")
    # Each object member ends with a two-space `},`; that lets us preserve multiline signatures.
    blocks = re.findall(r"^  (?:async\s+)?([A-Za-z_]\w*)\((.*?)\)\s*\{(.*?)^  \},", text, flags=re.M|re.S)
    rows=[]
    for name, signature, body in blocks:
        method = re.search(r'method:\s*"([A-Z]+)"', body)
        url = re.search(r'url:\s*([`\"])(.*?)\1', body, flags=re.S)
        response = re.search(r'return\s+call<(.+?)>\s*\(', body, flags=re.S)
        # axios shorthand is used by checkEmail and chatbot.
        if not method:
            axios = re.search(r"\.([gp]\w+)\s*<([^>]+)>\s*\(\s*([`\"'])(.*?)\3", body, flags=re.S)
            if axios:
                http = {"get":"GET", "post":"POST", "patch":"PATCH", "put":"PUT", "delete":"DELETE"}.get(axios.group(1), axios.group(1).upper())
                method_s, url_s, response_s = http, compact(axios.group(4)), compact(axios.group(2))
            else:
                method_s, url_s, response_s = "KHÔNG GỌI HTTP TRỰC TIẾP", "Xem ghi chú", "Dữ liệu được xử lý ở frontend"
        else:
            method_s = method.group(1)
            url_s = compact(url.group(2)) if url else "Không trích được URL"
            response_s = compact(response.group(1)) if response else "Kiểu trả về không khai báo qua call<T>"
        rows.append((name, compact(signature), method_s, url_s, response_s))
    return rows

def usages(api_obj, func, service_rel):
    token = api_obj + "." + func
    found=[]
    for p in ROOT.glob("src/**/*"):
        if p.suffix not in (".ts", ".tsx") or p.as_posix().endswith(service_rel):
            continue
        try:
            if token in p.read_text(encoding="utf-8"):
                found.append(p.relative_to(ROOT).as_posix())
        except UnicodeDecodeError:
            pass
    if not found:
        return "Chưa thấy page/component gọi trực tiếp trong src (có thể là API dự phòng, Admin/Risk hoặc chưa nối UI)."
    pages=[x for x in found if "/pages/" in x or x.startswith("src/pages/")]
    others=[x for x in found if x not in pages]
    result=pages+others
    return "; ".join(result[:5]) + (f"; +{len(result)-5} file" if len(result)>5 else "")

def input_text(sig, func):
    if func in ("listBusinessJobs", "getPortfolioByExpert"):
        return "Không có endpoint mới: hàm frontend gọi API khác/ lọc dữ liệu cục bộ. Tham số: " + sig
    return "Tham số hàm: " + (sig or "không có body/param")

def build():
    d=Document(); setup(d)
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(45)
    r=p.add_run("AITASKER"); r.bold=True; r.font.name="Arial"; r.font.size=Pt(14); r.font.color.rgb=RGBColor.from_string("2D6BA3")
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(16)
    r=p.add_run("CATALOG TOÀN BỘ API\nTHEO 5 NHÓM ĐƯỢC GIAO"); r.bold=True; r.font.name="Arial"; r.font.size=Pt(23); r.font.color.rgb=RGBColor.from_string("173F67")
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(16)
    r=p.add_run("Auth • Profile • Contract • Workspace • Common"); r.italic=True; r.font.name="Arial"; r.font.size=Pt(12)
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(38)
    r=p.add_run("Bản này KHÔNG lọc API theo page trong ảnh: ghi cả API đang có trong các service liên quan."); r.font.name="Arial"; r.font.size=Pt(10.5)
    d.add_page_break()
    all_parsed=[]
    for group, rel, api_obj in SOURCES:
        parsed=parse_methods(ROOT/rel); all_parsed.extend(parsed)
    direct=sum(1 for _,_,m,_,_ in all_parsed if m!="KHÔNG GỌI HTTP TRỰC TIẾP")
    head(d,"1. Cách đọc và phạm vi",1)
    para(d, f"Có {len(all_parsed)} hàm service. Trong đó {direct} hàm gọi HTTP trực tiếp; 2 hàm Profile chỉ ghép/lọc từ API khác tại frontend. Mỗi dòng dưới đây nói rõ: hàm nào, endpoint nào, tham số gì, kiểu dữ liệu trả về và file/page đang gọi.")
    box(d,"Điểm cần trả lời khi bảo vệ:", "Một hàm trong service chưa chắc đã có page đang dùng. Nếu cột Page/UI ghi ‘chưa thấy’, em nói đó là API đã được khai báo cho luồng khác (thường Risk/Admin/Marketplace) hoặc chưa nối UI, không được nói là page của mình đã chạy nó.")
    box(d,"Quy tắc xử lý chung:", "apiClient tự gắn Bearer token, thường unwrap response `{ success, message, data }` thành `data`. 401 xóa session/chuyển login; 403 là không đủ quyền. Loading/error/empty state cụ thể nằm tại page/component ở cột cuối.")
    table(d,["Nhóm service","Số hàm","Ghi chú"],[
        ["Auth",str(len(parse_methods(ROOT/'src/services/authServices.ts'))),"Login, registration, OTP, password và session."],
        ["Profile",str(len(parse_methods(ROOT/'src/services/profileService.ts'))),"KYB/KYC/portfolio và cả API Staff/Admin/proposal có trong file."],
        ["Contract + Workspace",str(len(parse_methods(ROOT/'src/services/contractService.ts'))),"Bao gồm Contract, milestone, termination, change request, review."],
        ["Common",str(len(parse_methods(ROOT/'src/services/notificationService.ts'))+len(parse_methods(ROOT/'src/services/chatbotService.ts'))),"Notification và chatbot; AppShell dùng authApi.me()."],
    ])
    d.add_page_break()
    head(d,"2. API hoạt động theo flow",1)
    para(d,"Cách đọc mũi tên: API phía trước trả về dữ liệu hoặc thay đổi trạng thái để UI mới được gọi API phía sau. Nếu điều kiện không đạt, flow dừng và UI hiển thị lỗi/thông báo thay vì gọi bước kế tiếp.")
    table(d,["Flow / người thực hiện","Thứ tự API gọi","Điều kiện chuyển bước và dữ liệu dùng"],[
        ["1. Đăng nhập email\nUser",
         "1) authApi.checkEmail(email)\n→ 2) authApi.login({ email, password })\n→ 3) authApi.me() (AppShell đồng bộ session)\n→ 4) authApi.refresh(refreshToken) khi cần làm mới token.",
         "checkEmail chỉ hỗ trợ kiểm tra UI. Login trả SessionUser (accessToken, refreshToken, role, accountStatus, fullName). saveSession lưu token; AppShell dùng me() để xác định role/layout. Nếu login/refresh trả 401: xoá session, về /login; không gọi me()."],
        ["2. Đăng ký thường + OTP\nBusiness hoặc Expert",
         "1) authApi.checkEmail(email)\n→ 2) authApi.sendOtp({ email })\n→ 3) authApi.verifyOtp({ email, otp })\n→ 4) authApi.register({ email, password, fullName, phone, role })\n→ 5) profile flow KYB hoặc KYC.",
         "Chỉ khi email chưa tồn tại và OTP đúng mới gọi register. sendOtp trả expiresIn để đếm ngược. register trả SessionUser và trạng thái thường là Pending; tùy role, người dùng tiếp tục Business KYB hoặc Expert KYC."],
        ["3. Đăng ký/đăng nhập Google\nBusiness hoặc Expert",
         "Google credential\n→ authApi.googleLogin({ credential, role? })\n→ nếu chưa có tài khoản/chưa đủ thông tin: authApi.googleSignup({ credential, fullName?, phone, role })\n→ authApi.me().",
         "Google login trả SessionUser nếu tài khoản tồn tại. Đăng ký Google cần phone + role; sau saveSession, điều hướng về trang profile tương ứng. Lỗi credential hoặc role không hợp lệ thì dừng tại Auth page."],
        ["4. Quên và đổi mật khẩu\nUser chưa đăng nhập",
         "1) authApi.checkEmail(email)\n→ 2) authApi.forgotPassword({ email })\n→ người dùng lấy token từ email/link\n→ 3) authApi.resetPassword({ token, newPassword })\n→ 4) chuyển /login.",
         "Chỉ gửi yêu cầu khi email hợp lệ. resetPassword chỉ gọi khi có token URL; token hết hạn/không hợp lệ: hiển thị lỗi, ở lại ResetPasswordPage."],
        ["5. Business KYB\nBusiness",
         "1) profileApi.getMyBusiness()\n→ (nhập taxCode) profileApi.checkTaxCode(taxCode)\n→ profileApi.uploadBusinessLicense(file)\n→ profileApi.upsertBusiness({ ...form, businessLicenseUrl })\n→ profileApi.getMyBusiness() để reload badge/form.",
         "Tax check trả thông tin doanh nghiệp để điền preview. Upload trả URL/path, URL này được đưa vào payload upsert. Nếu API lỗi: không gọi upsert khi chưa có file/URL cần thiết; UI giữ form và báo lỗi. upsert trả kybStatus/rejectionReason để hiển thị Pending/Approved/Rejected."],
        ["6. Expert KYC + Portfolio\nExpert",
         "1) profileApi.getMyExpert() + profileApi.getMyPortfolio()\n→ profileApi.uploadExpertPortfolio(file)\n→ profileApi.upsertExpert({ ...KYC, portfolioUrl })\n→ profileApi.uploadExpertCertificate(file) (lặp cho từng file)\n→ profileApi.upsertPortfolio({ domainIds, skillIds, technologyIds, certificates, ... })\n→ getMyExpert()/getMyPortfolio() reload.",
         "Upload trả URL/path; URL portfolio/chứng chỉ được nhét vào payload save. Danh mục domain/skill/technology được load ở page để chọn ID. KYC/portfolio lỗi thì giữ dữ liệu đang nhập, không chuyển trạng thái thành công."],
        ["7. Tạo và kích hoạt hợp đồng\nBusiness + Expert",
         "(Từ màn Manage Job) contractApi.createFromProposal(proposalId, payload)\n→ contractApi.getContract(id) / listMilestones(id)\n→ Business và Expert lần lượt contractApi.sign(id)\n→ cả hai lần lượt contractApi.signNda(id)\n→ Business contractApi.payDeposit(id)\n→ Expert contractApi.payExpertDeposit(id)\n→ getContract(id) reload trạng thái ACTIVE.",
         "Tạo từ proposal là luồng Marketplace nhưng API nằm trong contractService. Hợp đồng chỉ sang PENDING khi đủ hai chữ ký contract và hai NDA. Business deposit giữ 20% và Expert deposit giữ 10%; khi hai khoản HELD, backend kích hoạt ACTIVE. needTopup=true thì dừng ở deposit, yêu cầu nạp tiền."],
        ["8. Mở Workspace và làm milestone\nBusiness + Expert",
         "1) contractApi.getContract(contractId)\n→ 2) contractApi.listMilestones(contractId)\n→ khi chọn milestone: listCriteria(milestoneId) + listDeliverables(milestoneId) + listProgressReports(contractId, milestoneId)\n→ Business depositMilestoneEscrow(contractId, milestoneId)\n→ contractApi.startMilestone(milestoneId).",
         "WorkspacePage luôn tải contract/timeline trước. Business ký quỹ milestone để backend kiểm tra contract ACTIVE; sau đó mới start. start trả Milestone trạng thái IN_PROGRESS. Nếu chưa ký quỹ/contract chưa ACTIVE, backend từ chối và UI giữ trạng thái hiện tại."],
        ["9. Báo cáo tiến độ\nBusiness + Expert",
         "Business requestProgressReport(contractId, milestoneId)\n→ Expert submitProgressReport(contractId, milestoneId, payload)\n→ listProgressReports(...) reload lịch sử\n→ Business acknowledgeProgressReport(..., reportId) hoặc feedbackProgressReport(..., reportId, payload)\n→ Expert sửa/nộp report tiếp theo nếu cần.",
         "request trả requestNumber/status/dueAt. Report trả percent, checkpointType, links, isLate, acknowledgementState. acknowledge nghĩa là đã xem; feedback có feedback/category/severity/requiresAdjustment để Expert điều chỉnh."],
        ["10. Nộp và nghiệm thu deliverable\nExpert rồi Business",
         "Expert uploadMilestoneSourceCode(contractId, milestoneId, file)\n→ submitDeliverable(contractId, milestoneId, { sourceCodeUrl, submissionUrl, demoLink, notes })\n→ Business listDeliverables(milestoneId) / listCriteria(milestoneId)\n→ approveMilestone(...) hoặc rejectMilestone(..., { reason, failedCriteria })\n→ listMilestones(contractId) reload.",
         "Upload source trả URL; URL đi vào payload deliverable. approve trả Milestone COMPLETED và backend release 100% escrow milestone. reject trả milestone IN_PROGRESS, rejectCount/rejectionFeedback và deliverable REJECTED để Expert sửa rồi nộp lại."],
        ["11. Thay đổi hợp đồng\nBusiness hoặc Expert",
         "createChangeRequest(contractId, payload)\n→ listChangeRequests(contractId)\n→ phía còn lại acceptChangeRequest(contractId, requestId, reviewNote) hoặc rejectChangeRequest(...)",
         "create/accept/reject trả ContractChangeRequest. Sau accept, page reload contract/milestone nếu request làm thay đổi điều khoản. Đây là API Contract mở rộng; một số list page chưa gọi trực tiếp."],
        ["12. Chấm dứt / tranh chấp\nBusiness, Expert, Staff/Admin",
         "immediateTermination(...) hoặc requestTermination(contractId, payload)\n→ listTerminationRequests(contractId) / getTerminationRequest(id)\n→ acceptTerminationRequest(id) | disputeTerminationRequest(id, reason) | withdrawTerminationRequest(id, reason)\n→ Staff assignTerminationStaff(id, staffId)\n→ Admin approveTerminationRequest(id, payload)\n→ executeTerminationSettlement(id) hoặc refundTerminationDeposit(id, payload).",
         "Tùy trạng thái và quyền mà chỉ một nhánh được gọi. File chứng cứ: createCaseAttachment/listCaseAttachments; bằng chứng một phần: submitTerminationPartialEvidence. Các API Admin/Staff chưa chắc đã có UI trực tiếp trong mã hiện tại, catalog bên dưới ghi rõ."],
        ["13. Notification và chatbot\nMọi user đăng nhập",
         "AppShell: notificationApi.list() + notificationApi.unreadCount()\n→ NotificationPage: markRead(notificationId) hoặc markAllRead()\n→ UI reload list/badge.\nChatbot: chatbotApi.ask(question) → render ChatbotResponse.",
         "Notification list trả NotificationItem[]; unreadCount tạo badge. Mark API trả item đã đổi isRead. Chatbot chỉ gọi ask khi user gửi câu hỏi; response được render thành câu trả lời, lỗi thì ChatBox hiển thị fallback."],
    ])
    for idx,(group,rel,api_obj) in enumerate(SOURCES, start=3):
        d.add_page_break()
        parsed=parse_methods(ROOT/rel)
        head(d,f"{idx}. {group} - toàn bộ {len(parsed)} hàm",1)
        para(d,f"Source: {rel}. Cột Page/UI được dò trực tiếp từ việc tìm `{api_obj}.tenHam(...)` trong mã nguồn frontend.")
        rows=[]
        for n,(func,sig,http,url,response) in enumerate(parsed,1):
            feature=MEANING.get(func, func)
            endpoint=(http+" "+url) if http!="KHÔNG GỌI HTTP TRỰC TIẾP" else http
            response_show="Trả về: "+response
            rows.append([f"{n}. {feature}\n({func})",endpoint,input_text(sig,func)+"\n"+response_show,usages(api_obj,func,rel)])
        headers=["# / Chức năng (hàm)","HTTP endpoint","Input và kiểu trả về","Page / component đang gọi"]
        # Auth has only 11 rows, but the first 9 used to leave a nearly empty continuation page.
        # Split intentionally so both reference pages remain useful when printed.
        if group == "AUTH":
            table(d,headers,rows[:6])
            para(d,"AUTH - tiếp theo (các API đăng ký, password và Google):",lead="AUTH - tiếp theo")
            table(d,headers,rows[6:])
        elif group == "PROFILE":
            table(d,headers,rows[:10])
            para(d,"PROFILE - tiếp theo (API danh sách, Staff review và helper):",lead="PROFILE - tiếp theo")
            table(d,headers,rows[10:])
        else:
            table(d,headers,rows)
        if group=="CONTRACT + WORKSPACE":
            para(d,"Lưu ý phân nhóm: các hàm termination/dispute/refund thuộc luồng Risk; createFromProposal, Job milestone, change request/review thường thuộc Marketplace/Contract mở rộng. Chúng vẫn được liệt kê vì đều thực sự nằm trong contractService.ts.")
    d.add_page_break(); head(d,"8. Cách dùng tài liệu khi được hỏi",1)
    para(d,"Bước 1: tìm theo tên chức năng hoặc tên hàm. Bước 2: đọc endpoint và tham số. Bước 3: mở file ở cột Page/UI để xem chính xác UI dùng field trả về ra sao. Nếu chưa thấy Page/UI, trả lời trung thực rằng frontend hiện chưa gọi trực tiếp, thay vì tự nhận đã triển khai page đó.")
    para(d,"Ví dụ trả lời: ‘contractApi.approveMilestone gọi POST /api/v1/contracts/{contractId}/milestones/{milestoneId}/approve. Hàm trả về Milestone; WorkspacePage gọi nó để Business duyệt milestone rồi reload dữ liệu.’")
    d.save(OUT)
    print(OUT)

if __name__ == '__main__': build()
