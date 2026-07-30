from pathlib import Path
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SOURCE = Path(r"D:\FPT\SUMMER_2026\SWP391\PROJECT_GUIDE.md")
OUTPUT = Path(r"D:\FPT\SUMMER_2026\SWP391\PROJECT_GUIDE_formatted.docx")

NAVY = "163A5F"
BLUE = "1F5D8F"
LIGHT_BLUE = "EAF2F8"
PALE = "F7F9FB"
GRAY = "5B6573"

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)

def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    elem = OxmlElement('w:tblHeader')
    elem.set(qn('w:val'), 'true')
    trPr.append(elem)

def prevent_row_split(row):
    trPr = row._tr.get_or_add_trPr()
    elem = OxmlElement('w:cantSplit')
    trPr.append(elem)

def add_page_field(paragraph):
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    paragraph._p.append(begin)
    paragraph._p.append(instr)
    paragraph._p.append(end)

def set_cell_width(cell, twips):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(twips))
    tcW.set(qn('w:type'), 'dxa')

def set_table_geometry(table, widths):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.first_child_found_in('w:tblW')
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(sum(widths)))
    tblW.set(qn('w:type'), 'dxa')
    tblInd = tblPr.first_child_found_in('w:tblInd')
    if tblInd is None:
        tblInd = OxmlElement('w:tblInd')
        tblPr.append(tblInd)
    tblInd.set(qn('w:w'), '0')
    tblInd.set(qn('w:type'), 'dxa')
    grid = tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn('w:w'), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)

def clean_inline(text):
    return re.sub(r'`([^`]+)`', r'\1', text).replace('**', '').replace('__', '')

def add_rich_paragraph(doc, text, style=None, italic=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(5)
    text = text.replace('**', '').replace('__', '')
    chunks = re.split(r'(`[^`]+`)', text)
    for chunk in chunks:
        if not chunk:
            continue
        run = p.add_run(chunk[1:-1] if chunk.startswith('`') else chunk)
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')
        run.italic = italic or chunk.startswith('`')
        if chunk.startswith('`'):
            run.font.name = 'Courier New'
            run._element.rPr.rFonts.set(qn('w:ascii'), 'Courier New')
            run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Courier New')
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor.from_string(NAVY)
    return p

def create_restartable_numbering(doc):
    """Create a Word numbering instance that starts at 1 for this list."""
    numbering = doc.part.numbering_part.element
    nums = numbering.findall(qn('w:num'))
    next_id = max([int(n.get(qn('w:numId'))) for n in nums] + [0]) + 1
    abstract_id = '7'  # Built-in decimal list used by the List Number style.
    num = OxmlElement('w:num')
    num.set(qn('w:numId'), str(next_id))
    abstract = OxmlElement('w:abstractNumId')
    abstract.set(qn('w:val'), abstract_id)
    num.append(abstract)
    override = OxmlElement('w:lvlOverride')
    override.set(qn('w:ilvl'), '0')
    start = OxmlElement('w:startOverride')
    start.set(qn('w:val'), '1')
    override.append(start)
    num.append(override)
    numbering.append(num)
    return next_id

def add_numbered_item(doc, text, num_id):
    p = add_rich_paragraph(doc, text)
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.22)
    p.paragraph_format.right_indent = Inches(0.05)
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), str(num_id))
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)
    return p

def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith('|'):
        cells = [clean_inline(c.strip()) for c in lines[i].strip().strip('|').split('|')]
        if not all(re.fullmatch(r':?-{3,}:?', c.replace(' ', '')) for c in cells):
            rows.append(cells)
        i += 1
    return rows, i

def add_table(doc, data):
    if not data:
        return
    cols = max(len(r) for r in data)
    data = [r + [''] * (cols - len(r)) for r in data]
    table = doc.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'
    total = 9360
    if cols == 2:
        widths = [2800, 6560]
    elif cols == 3:
        widths = [1050, 3950, 4360]
    elif cols == 4:
        widths = [700, 2800, 2900, 2960]
    else:
        widths = [total // cols] * cols
        widths[-1] += total - sum(widths)
    set_table_geometry(table, widths)
    for row_idx, values in enumerate(data):
        row = table.rows[0] if row_idx == 0 else table.add_row()
        for col_idx, value in enumerate(values):
            cell = row.cells[col_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(value)
            run.font.name = 'Arial'
            run._element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
            run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')
            run.font.size = Pt(8.15 if cols >= 4 else 8.65)
            if row_idx == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                set_cell_shading(cell, NAVY)
            elif row_idx % 2 == 0:
                set_cell_shading(cell, PALE)
    set_repeat_table_header(table.rows[0])
    for row in table.rows:
        prevent_row_split(row)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def configure(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Arial'
    normal._element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12
    for name, size, color in [('Heading 1', 17, NAVY), ('Heading 2', 13, BLUE), ('Heading 3', 11, NAVY)]:
        st = styles[name]
        st.font.name = 'Arial'
        st._element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
        st._element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(15 if name == 'Heading 1' else 10)
        st.paragraph_format.space_after = Pt(6)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(2)
    hr = header.add_run('AITASKER  |  PROJECT GUIDE')
    hr.font.name = 'Arial'
    hr.font.size = Pt(8)
    hr.font.color.rgb = RGBColor.from_string(GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run('AITASKER  |  Tài liệu nội bộ  |  Trang ')
    r.font.name = 'Arial'
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(GRAY)
    add_page_field(footer)

def build():
    markdown = SOURCE.read_text(encoding='utf-8')
    lines = markdown.splitlines()
    doc = Document()
    configure(doc)

    # Cover/title block: editorial-cover inspired, compact for a reference guide.
    doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kr = kicker.add_run('AITASKER')
    kr.bold = True
    kr.font.name = 'Arial'
    kr.font.size = Pt(12)
    kr.font.color.rgb = RGBColor.from_string(BLUE)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(44)
    title.paragraph_format.space_after = Pt(10)
    tr = title.add_run('Tài liệu tổng quan, luồng nghiệp vụ\nvà bản đồ chức năng')
    tr.bold = True
    tr.font.name = 'Arial'
    tr._element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
    tr._element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')
    tr.font.size = Pt(27)
    tr.font.color.rgb = RGBColor.from_string(NAVY)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(40)
    sr = subtitle.add_run('Tài liệu tham chiếu cho thành viên dự án, demo và bàn giao')
    sr.italic = True
    sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor.from_string(GRAY)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run('Nguồn: PROJECT_GUIDE.md | Phiên bản chuyển đổi Word')
    mr.font.size = Pt(9)
    mr.font.color.rgb = RGBColor.from_string(GRAY)
    doc.add_page_break()

    toc_title = doc.add_paragraph('Nội dung', style='Heading 1')
    toc_title.paragraph_format.space_after = Pt(12)
    toc_items = [
        '1. Dự án là gì?',
        '2. Vai trò và quyền chính',
        '3. Luồng nghiệp vụ end-to-end',
        '3A. API được sử dụng trong từng luồng',
        '4. Bản đồ trang, chức năng và API',
        '5. Danh mục API theo module',
        '6. Kiến trúc frontend',
        '7. Dữ liệu và trạng thái cần nhớ khi trả lời',
        '8. Câu trả lời nhanh cho các câu hỏi thường gặp',
        '9. Cách chạy và cấu hình',
        '10. Tài liệu/mã nguồn nên mở khi cần đào sâu',
    ]
    for item in toc_items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.name = 'Arial'
        run.font.size = Pt(10.5)
    doc.add_page_break()

    i = 0
    in_code = False
    code_buffer = []
    active_numbering_id = None
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()
        if stripped.startswith('```'):
            active_numbering_id = None
            if in_code:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2)
                p.paragraph_format.space_after = Pt(8)
                r = p.add_run('\n'.join(code_buffer))
                r.font.name = 'Consolas'
                r._element.rPr.rFonts.set(qn('w:ascii'), 'Consolas')
                r._element.rPr.rFonts.set(qn('w:hAnsi'), 'Consolas')
                r.font.size = Pt(8.5)
                r.font.color.rgb = RGBColor.from_string(NAVY)
                code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buffer.append(raw)
            i += 1
            continue
        if stripped.startswith('|'):
            active_numbering_id = None
            data, i = parse_table(lines, i)
            add_table(doc, data)
            continue
        if not stripped or stripped == '---':
            active_numbering_id = None
            i += 1
            continue
        if stripped.startswith('>'):
            active_numbering_id = None
            p = add_rich_paragraph(doc, stripped.lstrip('> ').strip(), italic=True)
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.right_indent = Inches(0.25)
            i += 1
            continue
        match = re.match(r'^(#{1,3})\s+(.+)$', stripped)
        if match:
            active_numbering_id = None
            level, content = len(match.group(1)), clean_inline(match.group(2))
            if level == 1:
                # The main H1 is represented by the cover title.
                i += 1
                continue
            doc.add_paragraph(content, style=f'Heading {level}')
            i += 1
            continue
        if re.match(r'^\d+\.\s+', stripped):
            if active_numbering_id is None:
                active_numbering_id = create_restartable_numbering(doc)
            add_numbered_item(doc, re.sub(r'^\d+\.\s+', '', stripped), active_numbering_id)
            i += 1
            continue
        if stripped.startswith('- '):
            # A nested bullet can appear inside a numbered list; preserve the
            # current list id so the following numbered item resumes correctly.
            add_rich_paragraph(doc, stripped[2:], style='List Bullet')
            i += 1
            continue
        active_numbering_id = None
        add_rich_paragraph(doc, stripped)
        i += 1
    doc.core_properties.title = 'AITASKER – Tài liệu tổng quan, luồng nghiệp vụ và bản đồ chức năng'
    doc.core_properties.subject = 'Project guide'
    doc.core_properties.author = 'AITASKER Team'
    doc.save(OUTPUT)
    print(OUTPUT)

if __name__ == '__main__':
    build()
