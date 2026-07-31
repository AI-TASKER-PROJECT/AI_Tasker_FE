from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from create_business_handbook import setup, table, para, bullet, head, box

OUT=Path(r"D:\FPT\SUMMER_2026\SWP391\TU_DIEN_PAGE_API_AITASKER.docx")

def section(d, title, route, purpose):
    head(d,title,1); para(d,f'Route: {route}. ',lead=f'Route: {route}. '); para(d,purpose)

def build():
 d=Document(); setup(d)
 q=d.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; q.paragraph_format.space_before=Pt(45); r=q.add_run('AITASKER'); r.bold=True; r.font.name='Arial'; r.font.size=Pt(14); r.font.color.rgb=RGBColor.from_string('2D6BA3')
 q=d.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; q.paragraph_format.space_before=Pt(16); r=q.add_run('TỪ ĐIỂN PAGE – API\nCHI TIẾT'); r.bold=True; r.font.name='Arial'; r.font.size=Pt(25); r.font.color.rgb=RGBColor.from_string('173F67')
 q=d.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; q.paragraph_format.space_before=Pt(15); r=q.add_run('Dùng để trả lời: page này làm gì, gọi API nào, gửi gì, nhận gì, hiển thị gì và lỗi ra sao'); r.italic=True; r.font.name='Arial'; r.font.size=Pt(11)
 q=d.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; q.paragraph_format.space_before=Pt(38); r=q.add_run('Phạm vi: Auth • Profile/Verification • Contract/Workspace • Notification • Wallet/PayOS • Chatbot'); r.font.name='Arial'; r.font.size=Pt(10.5)
 d.add_page_break()

 head(d,'1. Cách đọc mọi API trong tài liệu',1)
 table(d,['Khái niệm','Nghĩa thực tế'],[
 ['Method','GET = lấy dữ liệu; POST = tạo/thực hiện hành động; PATCH/PUT = cập nhật; DELETE = xóa.'],
 ['Path parameter','Biến nằm trong URL, ví dụ :contractId. Nó chỉ rõ bản ghi nào đang được thao tác.'],
 ['Query parameter','Biến sau dấu ?, ví dụ ?status=Approved. Dùng cho lựa chọn/lọc nhỏ.'],
 ['Body','Dữ liệu form gửi trong JSON. Upload file dùng FormData thay vì JSON.'],
 ['Response chung','Hầu hết API backend bọc trong { success, message, data }. Hàm call<T>() ở frontend lấy riêng data.'],
 ['Nguồn quyền','Axios tự gắn Authorization: Bearer accessToken. Backend kiểm tra token, role và trạng thái; UI không có quyền tự quyết.'],
 ])
 box(d,'Công thức trả lời nhanh:', '“Ở page này, người dùng bấm …; frontend gọi METHOD URL với body/params …; backend kiểm tra …; data trả về gồm … nên UI hiển thị …; khi đang chờ thì …, khi rỗng thì …, khi lỗi thì ….”')
 head(d,'Các lỗi chung do apiClient xử lý',2)
 for x in ['Timeout 10 giây: báo yêu cầu quá thời gian chờ.', '401: clearSession(), lưu message rồi chuyển về /login; nguyên nhân có thể token hết hạn hoặc login ở nơi khác.', '403: message “không có quyền”.', 'Các code nghiệp vụ được dịch: INSUFFICIENT_BALANCE → không đủ số dư; CONTRACT_INVALID_STATUS → chưa đủ điều kiện ký quỹ; PROGRESS_REPORT_ACK_PENDING → report trước chưa được Business xác nhận.', 'FormData upload: interceptor bỏ Content-Type JSON để browser tự đặt multipart/form-data kèm boundary.']:
   bullet(d,x)

 section(d,'2. LoginPage – đăng nhập','/login','Mục đích: tạo session cho tài khoản hiện có hoặc xử lý luồng Google. Khi thành công, session được lưu và user được điều hướng vào app phù hợp role/status.')
 table(d,['Thao tác UI','API, input gửi','data nhận về','UI dùng data để làm gì?'],[
 ['Bấm Đăng nhập email/password','GET /api/auth/check-email?email=… trước để phát hiện email chưa tồn tại. Sau đó POST /api/auth/login body { email, password }.','SessionUser { accessToken, refreshToken, role, accountStatus, email, fullName }.','saveSession(); token đi kèm API sau; role hiển thị sidebar; accountStatus dùng profile gate.'],
 ['Chọn Google','POST /api/auth/google/login body { credential, role? }. Nếu là user Google mới, UI có bước chọn role, fullName/phone và POST /api/auth/google/register.','SessionUser giống login thường.','Lưu session/điều hướng. Nếu backend chưa cấu hình Google hoặc credential sai thì hiện message.'],
 ['Mở lại app/token cũ','GET /api/auth/me; khi cần POST /api/auth/refresh body { refreshToken }.','me trả role/accountStatus/email/fullName; refresh trả SessionUser/token mới.','Đồng bộ session, tránh dùng token hết hạn.'],
 ])
 para(d,'State thật trong page: loading=true trước request và Button nhận prop loading để khóa submit. catch hiển thị message. LoginPage còn đọc sessionStorage message khi interceptor chuyển về login.')

 section(d,'3. RegisterPage, ForgotPasswordPage, ResetPasswordPage','/register • /forgot-password • /reset-password','Ba page này tạo tài khoản có xác minh email và cho phép khôi phục mật khẩu an toàn.')
 table(d,['Page / bước','API, request','Kết quả/hiển thị','Khi lỗi hoặc chờ'],[
 ['/register – bước kiểm email','GET /api/auth/check-email?email=…','boolean: true nghĩa email đã tồn tại; UI chặn không gửi OTP.','loading; hiện message email tồn tại/không hợp lệ.'],
 ['/register – gửi OTP','POST /api/auth/email/send-otp body { email }.','{ expiresIn }: UI hiển thị/đếm thời gian OTP và mở bước nhập OTP.','Cho resend OTP; disable khi loading.'],
 ['/register – xác minh + tạo account','POST /api/auth/email/verify-otp body { email, otp }, sau đó POST /api/auth/register body { email,password,fullName,phone,role }.','register trả SessionUser. Account backend tạo status Pending; UI lưu session nhưng cần hoàn thiện KYC/KYB.','OTP sai/hết hạn báo lỗi. Form không bị mất dữ liệu. Role UI chỉ BUSINESS hoặc EXPERT.'],
 ['/forgot-password','GET check-email rồi POST /api/auth/forgot-password body { email }.','Không trả dữ liệu nghiệp vụ để hiện; UI báo đã gửi hướng dẫn email.','Email không tồn tại/lỗi backend hiện Notice.'],
 ['/reset-password','POST /api/auth/reset-password body { token, newPassword }. Token lấy từ URL.','Thành công: UI báo và chuyển /login.','Token hết hạn/không hợp lệ hoặc password không đạt validation: hiện error.'],
 ])

 section(d,'4. BusinessProfilePage và BusinessVerificationProfilePage','/app/business/profile • /app/business/kyb','Mục đích: Business xem/chỉnh hồ sơ và nộp KYB. Hai page cùng làm việc với business profile nhưng KYB tập trung vào tax code, giấy phép và trạng thái xác minh.')
 table(d,['Page / thao tác','API / request','Data trả về và phần UI hiển thị','State/UI'],[
 ['Profile load','GET /api/v1/profiles/business/me. Page còn gọi listBusinessJobs(businessId) để hiện jobs đã đăng.','BusinessProfile: businessId, taxCode, companyName, address, verifiedRepresentative, businessLicenseUrl, kybStatus, rejectionReason.','Đổ vào form, StatusBadge và phần giấy phép/jobs. Không có profile: EmptyState/CTA.'],
 ['Kiểm tra mã số thuế ở KYB','GET /api/auth/tax-check/:taxCode.','TaxCheckResponse { taxCode, companyName, address, representative, status }.','Preview thông tin doanh nghiệp trước khi gửi; taxPreviewLoading riêng.'],
 ['Upload giấy phép','POST /api/v1/profiles/business/license-file, FormData { file }.','string URL/path.','UI lưu URL vào businessLicenseUrl, preview file; file lỗi: Notice, giữ form.'],
 ['Lưu/nộp lại KYB','POST /api/v1/profiles/business body Partial<BusinessProfile>, quan trọng taxCode + businessLicenseUrl.','BusinessProfile mới nhất, gồm kybStatus/rejectionReason.','Cập nhật form/badge. Backend có thể đưa lại Pending để Staff duyệt lại; UI nói rõ đang chờ duyệt.'],
 ])
 box(d,'Backend kiểm tra:', 'Tax code bắt buộc, đúng 10/13 chữ số, không trùng account khác. Backend lấy tên công ty/địa chỉ/đại diện từ tax-check. Đây là lý do dữ liệu companyName không chỉ tin dữ liệu form.')

 section(d,'5. ExpertProfilePage, ExpertVerificationProfilePage, ExpertPortfolioPage','/app/expert/profile • /app/expert/kyc • /app/expert/portfolio','Mục đích: Expert chứng minh danh tính và năng lực. Profile là KYC cơ bản; Portfolio là năng lực, skill/domain/technology và chứng chỉ.')
 table(d,['Page / thao tác','API / request','Data/UI','State/UI'],[
 ['Expert Profile load','Promise.all: GET /api/v1/profiles/expert/me; GET /api/v1/profiles/portfolio/me; GET catalog domains/skills/technologies; GET contracts.','ExpertProfile { nationalId, portfolioUrl, yearsOfExperience, kycStatus, rejectionReason }; catalogs tạo các lựa chọn chip.','Đổ form, status, dữ liệu tham chiếu. Các API phụ có catch → null/[] để page vẫn mở.'],
 ['Upload portfolio file','POST /api/v1/profiles/expert/portfolio-file FormData { file }.','string URL/path.','URL gán portfolioUrl; hiển thị file đã chọn/đã lưu.'],
 ['Lưu KYC','POST /api/v1/profiles/expert body { nationalId, portfolioUrl, yearsOfExperience, ... }.','ExpertProfile mới nhất.','Backend yêu cầu 3 trường trên; nationalId không trùng. Submission có thể quay account về Pending.'],
 ['Portfolio load/lưu','GET /api/v1/profiles/portfolio/me; POST /api/v1/profiles/portfolio body { domainIds, skillIds, technologyIds, yearsExperience, certificates, selfDescription }.','Portfolio; UI hiển thị chip lĩnh vực/kỹ năng/công nghệ, mô tả, years experience.','Empty text cho catalog/chứng chỉ; error Notice; Button loading.'],
 ['Upload chứng chỉ','POST /api/v1/profiles/portfolio/certificate-file FormData { file }.','string URL/path.','Lưu trong certificates, preview/link chứng chỉ.'],
 ])

 section(d,'6. VerificationsPage và VerificationDetailPage','/app/verifications • /app/verifications/:type/:id','Mục đích: Staff duyệt KYB/KYC. Lưu ý route frontend cho Staff/Admin vào được, nhưng service backend approveProfile hiện requireRole("STAFF").')
 table(d,['Thao tác','API / input','Data/UI','Xử lý'],[
 ['Mở danh sách','GET /api/v1/profiles/business và GET /api/v1/profiles/expert.','BusinessProfile[]/ExpertProfile[]; UI gộp, lọc theo loại/status và link detail.','Danh sách rỗng: EmptyState. Hiện page code chưa có catch riêng ở load, nên lỗi request cần được kiểm tra khi demo.'],
 ['Mở chi tiết Business','GET /api/v1/profiles/business/:id; nếu có file: GET /api/v1/profiles/files/view-url?path=…','Profile + signed/view URL; UI hiện thông tin tax/license/status.','Không tìm thấy: EmptyState.'],
 ['Mở chi tiết Expert','GET /api/v1/profiles/expert/:id; GET /api/v1/profiles/portfolio; có thể GET file view URL.','Expert profile và portfolio tìm theo expertId; UI hiện KYC, kinh nghiệm, chứng chỉ.','Danh sách/portfolio lỗi thì fallback để không crash.'],
 ['Approve/Reject','POST /api/v1/profiles/approve/{BUSINESS|EXPERT}/{profileId}?status=Approved|Rejected&reason=…','BusinessProfile hoặc ExpertProfile đã cập nhật status/approvedBy/rejectionReason.','Sau action UI cập nhật profile. Rejected cần reason; backend chỉ review profile Pending.'],
 ])

 section(d,'7. ContractsPage – danh sách hợp đồng','/app/contracts','Mục đích: cho user xem mọi contract được backend cho phép xem, lọc theo trạng thái và đi tới detail/workspace.')
 table(d,['Thời điểm / API','Data nhận','UI hiển thị'],[
 ['Lúc load: GET /api/v1/contracts','Contract[]: contractId, jobId, businessId, expertId, totalBudget, timelineDays, status, progress, timestamps, contractMilestones.','Card/list: title, bên tham gia, budget, progress, StatusBadge, link /app/contracts/:id và workspace.'],
 ['Làm giàu thông tin trong page: mỗi contract GET /api/v1/contracts/:id/change-requests và GET /api/v1/contracts/:id/milestones.','ContractChangeRequest[] và ContractMilestone[].','Số/yêu cầu thay đổi, tóm tắt milestone/status; request lỗi dùng [] để danh sách vẫn hiện.'],
 ])
 para(d,'Loading list lỗi trong code hiện tại có catch(() => setContracts([])); vì vậy UI đi vào EmptyState “chưa có contract” thay vì hiển thị lỗi phân biệt. Đây là điểm cần nói trung thực nếu bị hỏi về UX.')

 section(d,'8. ContractDetailPage – ký, NDA, ký quỹ và thay đổi thỏa thuận','/app/contracts/:contractId','Mục đích: là trang thỏa thuận. Nó không nộp sản phẩm; nó hiển thị/sửa các việc trước khi hoặc song song khi thực thi contract.')
 table(d,['Nút/chức năng','API – input','Data trả về / UI cập nhật'],[
 ['Load detail','GET /api/v1/contracts/:id; sau đó GET job milestones, contract milestones, disputes, change requests, Business/Expert profile, wallet.','Contract + ContractMilestone[] + party profile + wallet; UI hiện điều khoản, chữ ký, NDA, tiền ký quỹ, mốc, dispute/change request.'],
 ['Ký contract','POST /api/v1/contracts/:id/sign, không body.','Contract với businessAcceptedAt/expertAcceptedAt/status. UI refresh detail, badge chữ ký.'],
 ['Ký NDA','POST /api/v1/contracts/:id/nda-sign.','Contract với businessNdaSignedAt/expertNdaSignedAt/status. Khi đủ 4 dấu mốc, status PENDING.'],
 ['Expert từ chối','POST /api/v1/contracts/:id/reject.','Contract CANCELLED; backend đưa job về OPEN. UI thông báo/refresh.'],
 ['Business hủy nháp','POST /api/v1/contracts/:id/cancel-draft.','Contract CANCELLED. Chỉ nháp chưa có bất kỳ chữ ký/NDA nào mới được hủy.'],
 ['Ký quỹ contract','Business: POST /deposit/pay; Expert: POST /expert-deposit/pay. Không body.','PaymentActionResponse<ContractDeposit> { completed, needTopup, currentBalance, requiredAmount, missingAmount, data }. UI refresh contract + wallet; needTopup mở/hướng dẫn nạp tiền.'],
 ['Tạo change request','POST /api/v1/contracts/:id/change-requests body { changeType, changeSummary, proposedBudget?, proposedTimelineDays?, proposedScope?, proposedMilestones? }.','ContractChangeRequest; UI thêm vào lịch sử chờ review.'],
 ['Accept/Reject change','POST .../change-requests/:requestId/accept hoặc /reject body { reviewNote }.','Request có status/reviewNote. Accept thì page tải lại contract và milestones để hiện scope/budget đã đổi.'],
 ])
 bullet(d,'Mọi action có state riêng: ndaSubmitting, rejectLoading, cancelDraftLoading, changeRequestLoading, depositLoading. Nhờ vậy người dùng không gửi trùng.')

 section(d,'9. WorkspacePage – page có nhiều API nhất','/app/contracts/:contractId/workspace','Mục đích: thực thi milestone. Page load contract, milestone, dispute và termination request; khi chọn mốc sẽ load criteria, deliverables, reports để hiện tiến độ thực tế.')
 table(d,['Nút/chức năng','Endpoint + request','Data trả về để hiển thị'],[
 ['Tải workspace','GET /api/v1/contracts/:id; GET /api/v1/contracts/:id/milestones; GET /api/v1/contracts/:id/disputes; GET /api/v1/contracts/:id/termination-requests.','Contract, ContractMilestoneView[], Dispute[], TerminationRequest[]. UI hiện overview/timeline/tab trạng thái.'],
 ['Chi tiết mốc','GET /api/v1/milestones/:milestoneId/criteria; GET /api/v1/milestones/:milestoneId/deliverables; GET /api/v1/contracts/:id/milestones/:milestoneId/progress-reports.','AcceptanceCriteria[], Deliverable[], MilestoneProgressReport[]. UI hiện checklist, file/link, round nộp, report/feedback.'],
 ['Business ký quỹ mốc','POST /api/v1/contracts/:contractId/milestones/:milestoneId/deposit.','Milestone status mới. UI reload milestone; hiển thị đang thực hiện/escrow. Backend kiểm tra thứ tự và available balance.'],
 ['Expert start','POST /api/v1/milestones/:milestoneId/start.','Milestone. UI chuyển timeline sang In Progress nếu không đã do deposit auto-start.'],
 ['Business yêu cầu report','POST /api/v1/contracts/:contractId/milestones/:milestoneId/progress-report-request.','ProgressReportRequestRecord { requestNumber, status, dueAt }. UI badge yêu cầu/chờ nộp.'],
 ['Expert nộp report','POST /api/v1/contracts/:contractId/milestones/:milestoneId/progress-reports body { content, percentComplete?, attachmentUrl?, sourceCodeUrl?, sourceCodeFileUrl?, demoLink?, submissionNotes? }.','MilestoneProgressReport: checkpointType, isLate, acknowledgementState, content/links. UI thêm vào lịch sử.'],
 ['Business acknowledge/feedback','POST .../progress-reports/:reportId/acknowledge; POST .../feedback body { feedback, category?, severity?, dodItems?, requiresAdjustment? }.','Report đã ACKNOWLEDGED hoặc có businessFeedback/requiresAdjustment. UI hiện feedback; Expert mới có thể nộp report tiếp.'],
 ['Upload source ZIP','POST /api/v1/contracts/:contractId/milestones/:milestoneId/source-code-file FormData { file }.','string URL. UI dùng URL trong report/deliverable. Backend giới hạn ZIP tối đa 50MB.'],
 ['Nộp deliverable','POST /api/v1/contracts/:contractId/milestones/:milestoneId/deliverables body { sourceCodeUrl?, sourceCodeFileUrl?, demoLink?, submissionNotes? }.','Deliverable { deliverableId, submissionRound, status=SUBMITTED, links, notes }. UI hiện chờ review.'],
 ['Approve','POST /api/v1/contracts/:contractId/milestones/:milestoneId/approve.','Milestone COMPLETED, escrowReleasedAt. UI reload contract/milestones; backend chuyển tiền escrow 100% cho Expert.'],
 ['Reject','POST .../reject body { reason, failedCriteria?: [{ criteriaId, reason }] }.','Milestone quay IN_PROGRESS; rejectCount/resubmitCount/lastRejectionFeedback. UI hiển thị lý do & tiêu chí fail.'],
 ])
 table(d,['Các API rủi ro trên Workspace','Mục đích'],[
 ['POST .../milestones/check-overdue (Admin)','Đánh dấu overdue theo deadline, UI hiện cảnh báo/chặn nộp sản phẩm quá hạn.'],
 ['POST .../milestones/sla-auto-approve (Admin)','Duyệt tự động mốc UNDER_REVIEW quá SLA nếu không dispute/termination.'],
 ['POST .../disputes, POST /api/v1/disputes/:id/escalation-request','Tạo tranh chấp/đẩy Staff can thiệp; UI đóng băng luồng thường.'],
 ['POST termination request/accept/dispute/immediate-termination','Tạo/phản hồi/chống lại yêu cầu chấm dứt; UI theo termination status.'],
 ])

 section(d,'10. NotificationPage, WalletPage, PayOS và Chatbot','/app/notifications • /app/wallet • AppShell Chatbox','Các page/chức năng chung giúp user nhìn sự kiện, số tiền và hỗ trợ, nhưng chúng không tự quyết định nghiệp vụ contract.')
 table(d,['Page/chức năng','API','Data/UI/state'],[
 ['NotificationsPage load','GET /api/v1/notifications; GET /api/v1/notifications/unread-count.','NotificationItem[] { type,title,message,targetUrl,isRead,createdAt,metadata }; badge count. loading=true → spinner; [] → EmptyState.'],
 ['Đọc thông báo','PATCH /api/v1/notifications/:id/read; PATCH /api/v1/notifications/read-all.','NotificationItem/array mới; UI cập nhật isRead. Lỗi thì gọi refresh để đồng bộ lại.'],
 ['WalletPage load','GET /api/v1/wallet/me; GET wallet transactions; GET withdrawal requests.','SystemWallet { currentBalance, availableBalance, escrowBalance, holding/disputed balance, currency }; lịch sử để giải thích tiền đang ở đâu. Promise.allSettled để phần lỗi không làm hỏng toàn page.'],
 ['Nạp qua PayOS','POST /api/payments/payos/create body { amount, description }; POST /api/payments/payos/:orderCode/sync.','CreatePayOSPaymentResponse (checkout URL/QR/orderCode tùy type); sync trả PaymentOrder status. Chỉ khi backend xác nhận mới reload wallet.'],
 ['Chatbot','POST /api/chatbot/ask body { question }.','ChatbotResponse { answer, sources }. UI append câu hỏi/câu trả lời; loading khóa gửi; lỗi có fallback.'],
 ])

 head(d,'11. Phụ lục: bản đồ endpoint theo page',1)
 table(d,['Nhóm','GET (lấy/hiển thị)','POST/PATCH (hành động)'],[
 ['Auth','/api/auth/me; /api/auth/check-email','/login, /refresh, /register, /email/send-otp, /email/verify-otp, /forgot-password, /reset-password, /google/login, /google/register'],
 ['Profile','/profiles/business/me, /expert/me, /portfolio/me, /profiles/business, /expert, /files/view-url; /auth/tax-check/:taxCode','/profiles/business, /expert, /portfolio; /business/license-file, /expert/portfolio-file, /portfolio/certificate-file; /approve/:type/:id'],
 ['Contract','/contracts, /contracts/:id, /contracts/:id/milestones, /change-requests','/from-proposals/:proposalId, /sign, /nda-sign, /reject, /cancel-draft, /deposit/pay, /expert-deposit/pay, change accept/reject'],
 ['Workspace','/milestones/:id/criteria, /deliverables, /progress-reports; /contracts/:id/disputes','deposit, start, progress request/report/ack/feedback, source-code-file, deliverables, approve/reject, overdue/SLA, disputes, termination'],
 ['Common','/notifications, /unread-count, /wallet/me','notification read/read-all, PayOS create/sync, chatbot ask'],
 ])
 head(d,'12. Khi bạn cần giải thích 1 field response',1)
 for x in ['Contract.status: quyết định nút nào xuất hiện. DRAFT chỉ ký/NDA; PENDING ký quỹ; ACTIVE mở thực thi.', 'ContractMilestone.finalBudget: số tiền Business phải giữ trong escrow cho mốc đó. contractMilestone.status: quyết định action của mốc.', 'PaymentActionResponse.needTopup/missingAmount: quyết định hiện modal/link nạp tiền thay vì báo “thành công” giả.', 'Deliverable.submissionRound/status/rejectionFeedback: cho biết sản phẩm nộp lần mấy, có bị từ chối không và phải sửa gì.', 'MilestoneProgressReport.acknowledgementState/isLate/businessFeedback: cho biết Business đã xem chưa, report trễ không và có cần chỉnh sửa không.', 'NotificationItem.targetUrl: giúp UI đưa user từ notification đến đúng contract/dispute/page liên quan.']:
   bullet(d,x)
 d.core_properties.title='AITASKER - Từ điển Page API chi tiết'
 d.core_properties.author='AITASKER Team'
 d.save(OUT); print(OUT)

if __name__=='__main__': build()
